import os
import re
import json
import argparse
from pathlib import Path
from typing import List, Dict

SAMPLE_FILE_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/豫洛阳-兰郑长干线-CPY-0790-BFGDGS-ZZSYQFGS.docx"
SAMPLE_OUT_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/luxian_extract_result.txt"

# 统一别名（与 content_consistency_check.py 对齐）
EVAC_ALIASES = [
    "逃生路线图",
    "逃生路线",
]

ASSEMBLY_ALIASES = [
    "应急疏散集合点位置",
    "应急疏散集结点位置",
    "应急疏散集结点",
    "应急疏散集合点",
    "疏散集结点",
]

WINDOW = 2

try:
    from docx import Document
except ImportError:
    raise ImportError("请先 pip install python-docx")


def parse_args():
    ap = argparse.ArgumentParser(description="统一提取：逃生路线图 / 应急疏散集合点位置")
    ap.add_argument("--task", choices=["evac_route", "assembly_point"], default="evac_route",
                    help="选择任务：evac_route=逃生路线图；assembly_point=应急疏散集合点位置")
    ap.add_argument("--input", default=None, help="输入 DOCX 文件路径；不提供则优先使用样例，否则自动查找")
    ap.add_argument("--output_dir", default=None, help="图片输出目录；默认按任务自动命名")
    ap.add_argument("--out_path", default=None, help="结果写入该文本/JSON 文件（可选）")
    ap.add_argument("--first_only", action="store_true", help="仅提取第一张图片（输出单对象），否则输出列表")
    return ap.parse_args()


def _is_valid_docx_path(p: Path) -> bool:
    return p.suffix.lower() == ".docx" and not p.name.startswith("~$") and p.exists()


def _auto_discover_docx() -> Path | None:
    cwd = Path.cwd()
    candidates = []
    wm = cwd / "word-master"
    if wm.exists():
        for p in wm.rglob("*.docx"):
            if _is_valid_docx_path(p):
                candidates.append(p)
    for p in cwd.glob("*.docx"):
        if _is_valid_docx_path(p):
            candidates.append(p)
    if not candidates:
        for p in cwd.rglob("*.docx"):
            if _is_valid_docx_path(p):
                candidates.append(p)
                break
    return candidates[0] if candidates else None


def _save_images_from_runs(doc, runs, out_dir: Path, basename: str, counter: list) -> list:
    out = []
    for run in runs:
        drawing = run._element.xpath('.//a:blip')
        for blip in drawing:
            rEmbed = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rEmbed and rEmbed in doc.part.rels:
                rel = doc.part.rels[rEmbed]
                if "image" in rel.target_ref:
                    counter[0] += 1
                    image_data = rel.target_part.blob
                    img_path = out_dir / f"{basename}_{counter[0]}.png"
                    with open(img_path, 'wb') as f:
                        f.write(image_data)
                    out.append({"image_path": str(img_path), "index": counter[0]})
    return out


def extract_images_between_numbered_captions(doc_path: str,
                                             start_no: int,
                                             start_aliases: list[str],
                                             end_no: int,
                                             end_aliases: list[str],
                                             output_dir: str,
                                             out_name_prefix: str) -> list:
    """
    识别：(start_no) 含 start_aliases 文本  到  (end_no) 含 end_aliases 文本 之间的所有图片。
    返回保存路径列表。
    """
    doc = Document(doc_path)
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    doc_name = Path(doc_path).stem
    counter = [0]

    def numbered_re(n: int, aliases: list[str]):
        alias_pat = "|".join(re.escape(a) for a in aliases)
        return re.compile(rf'^\s*[（(]?\s*{n}\s*[)）].*?(?:{alias_pat})')

    start_re = numbered_re(start_no, start_aliases)
    end_re = numbered_re(end_no, end_aliases)

    in_range = False
    ended = False
    results = []

    for para in doc.paragraphs:
        if ended:
            break
        text = para.text or ""
        if not in_range:
            if start_re.search(text):
                in_range = True
            continue
        if end_re.search(text):
            ended = True
            break
        results.extend(_save_images_from_runs(doc, para.runs, out_dir, f"{doc_name}_{out_name_prefix}", counter))

    if not ended:
        for table in doc.tables:
            if ended:
                break
            for row in table.rows:
                if ended:
                    break
                for cell in row.cells:
                    if ended:
                        break
                    cell_text = cell.text or ""
                    if not in_range:
                        if start_re.search(cell_text):
                            in_range = True
                        continue
                    if end_re.search(cell_text):
                        ended = True
                        break
                    for p in cell.paragraphs:
                        results.extend(_save_images_from_runs(doc, p.runs, out_dir, f"{doc_name}_{out_name_prefix}", counter))

    return results


def extract_images_after_caption_until_next_number(doc_path: str,
                                                  aliases: list[str],
                                                  output_dir: str,
                                                  out_name_prefix: str) -> list:
    """
    在命中“(n) + 任一别名”的段落/单元格后，收集直到下一编号项前的所有图片。
    """
    doc = Document(doc_path)
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    doc_name = Path(doc_path).stem
    counter = [0]

    NUM_ITEM_PAT = re.compile(r'^\s*[（(]?\s*(\d+)\s*[)）]')
    alias_pat = "|".join(re.escape(a) for a in aliases)
    start_re = re.compile(rf'^\s*[（(]?\s*\d+\s*[)）].*?(?:{alias_pat})')

    in_range = False
    start_no = None
    results = []

    def is_next_numbered(text: str) -> bool:
        m = NUM_ITEM_PAT.match(text or "")
        if not m or start_no is None:
            return False
        try:
            num = int(m.group(1))
            return num != start_no
        except Exception:
            return False

    for para in doc.paragraphs:
        text = para.text or ""
        if not in_range:
            if start_re.search(text):
                in_range = True
                m = NUM_ITEM_PAT.match(text)
                start_no = int(m.group(1)) if m else None
                continue
            else:
                continue
        if is_next_numbered(text):
            break
        got = _save_images_from_runs(doc, para.runs, out_dir, f"{doc_name}_{out_name_prefix}", counter)
        if got:
            results.extend(got)

    if not results:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text or ""
                    if not in_range:
                        if start_re.search(cell_text):
                            in_range = True
                            m = NUM_ITEM_PAT.match(cell_text)
                            start_no = int(m.group(1)) if m else None
                        else:
                            continue
                    else:
                        if is_next_numbered(cell_text):
                            return results
                    for p in cell.paragraphs:
                        got = _save_images_from_runs(doc, p.runs, out_dir, f"{doc_name}_{out_name_prefix}", counter)
                        if got:
                            results.extend(got)
                    # 不在此处提前返回，继续扫描该单元格直至遇到下一编号项

    return results


def get_hca_types_from_docx(docx_path: str) -> list:
    """
    读取文档并解析“高后果区类型”，返回类型列表（去空、去重、已规范化）。
    """
    try:
        doc = Document(docx_path)
    except Exception:
        return []
    try:
        hca_type = extract_hca_type_from_tables_luxian(doc)
    except Exception:
        hca_type = None
    if not hca_type:
        return []
    # 使用与内容校验一致的拆分与规范
    s = _norm_value_luxian(hca_type)
    parts = [p for p in re.split(r"[、/|，,；; ]+", s) if p]
    seen, out = set(), []
    for p in parts:
        if p not in seen:
            seen.add(p)
            out.append(p)
    return out


def build_vlm_messages_from_docx(docx_path: str,
                                 prompt_text: str = (
                                     "请基于以下两类图片进行核查：‘(2) 逃生路线图’与‘应急疏散集结点位置’。\n"
                                     "请按如下要求作答：\n"
                                     "1) 逃生路线图：是否清晰标注了疏散路线、关键转折/出口、方向箭头；疏散方向是否指向管道两侧（远离管道）。\n"
                                     "2) 集结点位置：是否给出明显位置标注及文字说明；位置是否位于潜在影响半径范围以外。\n"
                                     "3) 请先给出总结论‘是否符合要求：是/否’，如不符合，逐项指出缺失/问题并给出具体改进建议（应补充哪些标注/如何调整路线或位置、示例文本）。"
                                 ),
                                 evac_output_dir: str | None = None,
                                 assembly_output_dir: str | None = None,
                                 first_only: bool = False) -> List[Dict]:
    """
    构建用于 VLM 的消息数组：优先提取‘(2) 逃生路线图’及‘应急疏散集结点位置’对应图片，并附上核查提示。

    返回格式：
    [
      {
        "role": "user",
        "content": [
          {"type": "image", "image": "...evac..."},
          {"type": "image", "image": "...assembly..."},
          {"type": "text",  "text":  "...提示词..."}
        ]
      }
    ]
    """
    docx_path = str(docx_path)
    # 人员密集型前置规则：直接文本结论，跳过模型
    types = get_hca_types_from_docx(docx_path)
    if "人员密集型" in types:
        return [{
            "role": "user",
            "content": [
                {"type": "text", "text": "符合要求：是（人员密集型前置规则，直接判定，无需调用模型）"}
            ]
        }]
    evac_out_dir = evac_output_dir or "逃生路线图输出"
    assembly_out_dir = assembly_output_dir or "应急疏散集合点输出"

    # 提取两类图片
    evac_items = extract_images_between_numbered_captions(
        docx_path, 2, EVAC_ALIASES, 3, ASSEMBLY_ALIASES, evac_out_dir, "逃生路线图"
    ) or []
    assembly_items = extract_images_after_caption_until_next_number(
        docx_path, ASSEMBLY_ALIASES, assembly_out_dir, "应急疏散集合点"
    ) or []

    if first_only:
        evac_items = evac_items[:1]
        assembly_items = assembly_items[:1]

    content: List[Dict] = []
    for it in evac_items:
        img_path = (it.get("image_path") if isinstance(it, dict) else it)
        if img_path:
            content.append({"type": "image", "image": img_path})
    for it in assembly_items:
        img_path = (it.get("image_path") if isinstance(it, dict) else it)
        if img_path:
            content.append({"type": "image", "image": img_path})

    content.append({"type": "text", "text": prompt_text})

    return [
        {
            "role": "user",
            "content": content,
        }
    ]


def build_vlm_messages_for_evac_route(docx_path: str,
                                      prompt_text: str = (
                                          "请核查‘(2) 逃生路线图’：\n"
                                          "1) 是否清晰标注疏散路线、关键转折/出口、方向箭头；\n"
                                          "2) 疏散方向是否指向管道两侧（远离管道）；\n"
                                          "请先给出‘是否符合要求：是/否’，如不符合请指出问题并给出改进建议。"
                                      ),
                                      output_dir: str | None = None,
                                      first_only: bool = False) -> List[Dict]:
    docx_path = str(docx_path)
    # 人员密集型前置规则：直接文本结论
    types = get_hca_types_from_docx(docx_path)
    if "人员密集型" in types:
        return [{
            "role": "user",
            "content": [
                {"type": "text", "text": "符合要求：是（人员密集型前置规则，直接判定，无需调用模型）"}
            ]
        }]
    out_dir = output_dir or "逃生路线图输出"
    evac_items = extract_images_between_numbered_captions(
        docx_path, 2, EVAC_ALIASES, 3, ASSEMBLY_ALIASES, out_dir, "逃生路线图"
    ) or []
    if first_only:
        evac_items = evac_items[:1]
    content: List[Dict] = []
    for it in evac_items:
        img_path = (it.get("image_path") if isinstance(it, dict) else it)
        if img_path:
            content.append({"type": "image", "image": img_path})
    content.append({"type": "text", "text": prompt_text})
    return [{"role": "user", "content": content}]


def build_vlm_messages_for_assembly_point(docx_path: str,
                                          prompt_text: str = (
                                              "请核查‘应急疏散集结点位置’：\n"
                                              "1) 是否给出明显位置标注及文字说明；\n"
                                              "2) 集结点位置是否位于潜在影响半径范围以外；\n"
                                              "请先给出‘是否符合要求：是/否’，如不符合请指出问题并给出改进建议。"
                                          ),
                                          output_dir: str | None = None,
                                          first_only: bool = False) -> List[Dict]:
    docx_path = str(docx_path)
    # 人员密集型前置规则：直接文本结论
    types = get_hca_types_from_docx(docx_path)
    if "人员密集型" in types:
        return [{
            "role": "user",
            "content": [
                {"type": "text", "text": "符合要求：是（人员密集型前置规则，直接判定，无需调用模型）"}
            ]
        }]
    out_dir = output_dir or "应急疏散集合点输出"
    assembly_items = extract_images_after_caption_until_next_number(
        docx_path, ASSEMBLY_ALIASES, out_dir, "应急疏散集合点"
    ) or []
    if first_only:
        assembly_items = assembly_items[:1]
    content: List[Dict] = []
    for it in assembly_items:
        img_path = (it.get("image_path") if isinstance(it, dict) else it)
        if img_path:
            content.append({"type": "image", "image": img_path})
    content.append({"type": "text", "text": prompt_text})
    return [{"role": "user", "content": content}]

# === 从 content_consistency_check.py 复制基础类型提取函数 ===
def _clean_luxian(s: str) -> str:
    return (s or "").strip()

def _norm_value_luxian(s: str) -> str:
    s = (s or "").strip()
    s = s.replace("类", "型")
    s = re.sub(r"[ \t]+", "", s)
    s = re.sub(r"[，,;/|]+", "、", s)
    s = re.sub(r"[：:；;。]$", "", s)
    return s

def extract_hca_type_from_tables_luxian(doc):
    """
    在表格中查找"高后果区类型"，与 content_consistency_check.py 保持一致
    """
    for tbl in getattr(doc, "tables", []):
        rows = list(getattr(tbl, "rows", []))
        for r_idx, row in enumerate(rows):
            cells = list(getattr(row, "cells", []))
            for c_idx, cell in enumerate(cells):
                raw = _clean_luxian(cell.text)
                if "高后果区类型" in raw.replace(" ", ""):
                    # 情况 A: 同单元格
                    m = re.search(r"高后果区类型[:：]?\s*(.+)", raw)
                    if m and _clean_luxian(m.group(1)):
                        return _norm_value_luxian(m.group(1))
                    # 情况 B: 同行右侧
                    right_texts = [_clean_luxian(c.text) for c in cells[c_idx + 1:] if _clean_luxian(c.text)]
                    seen, unique_texts = set(), []
                    for txt in right_texts:
                        if txt not in seen:
                            seen.add(txt)
                            unique_texts.append(txt)
                    right_val = "、".join(unique_texts).strip()
                    # 情况 C: 同列下几行
                    down_val = ""
                    if len(right_val) <= 1:
                        frags = []
                        for k in range(1, 4):
                            rr = r_idx + k
                            if rr < len(rows):
                                c2 = rows[rr].cells
                                if c_idx < len(c2):
                                    frags.append(_clean_luxian(c2[c_idx].text))
                        seen, unique_frags = set(), []
                        for f in frags:
                            if f and f not in seen:
                                seen.add(f)
                                unique_frags.append(f)
                        down_val = "、".join(unique_frags).strip()
                    val = right_val if len(right_val) > len(down_val) else down_val
                    val = _norm_value_luxian(val)
                    if val:
                        return val
    return None


def main():
    args = parse_args()
    chosen = None
    used_sample = False
    if args.input:
        chosen = Path(args.input)
    else:
        if SAMPLE_FILE_PATH and _is_valid_docx_path(Path(SAMPLE_FILE_PATH)):
            chosen = Path(SAMPLE_FILE_PATH)
            used_sample = True
        else:
            chosen = _auto_discover_docx()
    if not chosen or not Path(chosen).exists():
        print("未提供 --input，且未在当前目录或样例路径找到有效 DOCX（自动忽略临时锁文件 ~$.docx）。请使用 --input 指定文件，例如：")
        print("python luxian_extract.py --task evac_route --input C:/path/to/file.docx")
        return

    # --- 新增：类型识别及人员密集型判断 ---
    try:
        docx = Document(str(chosen))
    except Exception as e:
        print(f"读取文档失败: {e}")
        return
    hca_type = extract_hca_type_from_tables_luxian(docx)
    print(f"高后果区类型: {hca_type if hca_type else '未识别'}")
    type_list = []
    if hca_type:
        type_list = [t for t in re.split(r"[、/|，,；; ]+", hca_type) if t]
    if "人员密集型" not in type_list:
        print("【跳过】当前文档高后果区类型不是人员密集型，未提取图片。")
        return

    # 逃生路线图（编号2）输出目录
    evac_out_dir = args.output_dir or "逃生路线图输出"
    evac_imgs = extract_images_between_numbered_captions(
        str(chosen), 2, EVAC_ALIASES, 3, ASSEMBLY_ALIASES, evac_out_dir, "逃生路线图"
    )
    evac_img_path = evac_imgs[0]["image_path"] if evac_imgs else None

    # 应急疏散集合点输出目录
    assembly_out_dir = args.output_dir or "应急疏散集合点输出"
    assembly_imgs = extract_images_after_caption_until_next_number(
        str(chosen), ASSEMBLY_ALIASES, assembly_out_dir, "应急疏散集合点"
    )
    assembly_img_path = assembly_imgs[0]["image_path"] if assembly_imgs else None

    result = {
        "evac_route_image_path": evac_img_path,
        "assembly_point_image_path": assembly_img_path
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))

    out_file = args.out_path if args.out_path else (SAMPLE_OUT_PATH if used_sample else None)
    if out_file:
        try:
            Path(out_file).parent.mkdir(parents=True, exist_ok=True)
            with open(out_file, "w", encoding="utf-8") as f:
                f.write(json.dumps(result, ensure_ascii=False, indent=2))
        except Exception as e:
            print(f"写入结果失败: {e}")


if __name__ == "__main__":
    main()


