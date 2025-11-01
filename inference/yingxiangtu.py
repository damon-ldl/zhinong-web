from pathlib import Path
import argparse
import json
import re

# 样例文件与输出（参考 main.py）
SAMPLE_FILE_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/豫洛阳-兰郑长干线-CPY-0790-BFGDGS-ZZSYQFGS.docx"
SAMPLE_OUT_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/content_consistency_check.txt"

# 移除冗余的提取函数，统一通过区间提取


def extract_images_between_headings(docx_path,
                                   start_keyword="一、高后果区基本信息",
                                   end_heading_regex=r"二[、,.．\.]",
                                   output_dir="影像图输出"):
    """
    提取位于 start_keyword 之后，且在下一个符合 end_heading_regex 的标题出现之前的所有图片（正文与表格）。
    - start_keyword: 起始分界文本（包含该行之后开始计入）
    - end_heading_regex: 结束标题正则（匹配到即停止，常见为“二、”等）
    返回: [{image_path, index}]，未命中返回空列表。
    """
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
    try:
        doc = Document(docx_path)
    except PackageNotFoundError:
        return []
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    doc_name = Path(docx_path).stem
    in_range = False
    ended = False
    img_idx = 0
    results = []
    end_prog = re.compile(end_heading_regex)

    def save_from_runs_if_in_range(runs):
        nonlocal img_idx
        local = []
        for run in runs:
            drawing = run._element.xpath('.//a:blip')
            for blip in drawing:
                rEmbed = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rEmbed and rEmbed in doc.part.rels:
                    rel = doc.part.rels[rEmbed]
                    if "image" in rel.target_ref:
                        img_idx += 1
                        image_data = rel.target_part.blob
                        img_filename = f"{doc_name}_影像图_{img_idx}.png"
                        img_path = output_dir / img_filename
                        with open(img_path, 'wb') as f:
                            f.write(image_data)
                        local.append({"image_path": str(img_path), "index": img_idx})
        return local

    # 扫描正文
    for para in doc.paragraphs:
        if ended:
            break
        text = para.text or ""
        if not in_range:
            if start_keyword in text:
                in_range = True
            continue
        # 已进入范围后，先检查是否遇到结束标题
        if end_prog.search(text):
            ended = True
            break
        results.extend(save_from_runs_if_in_range(para.runs))

    # 扫描表格
    if not ended:
        for table in doc.tables:
            if ended:
                break
            for row in table.rows:
                if ended:
                    break
                for cell in row.cells:
                    if ended:
                        break
                    cell_text = cell.text or ""
                    if not in_range:
                        if start_keyword in cell_text:
                            in_range = True
                        continue
                    if end_prog.search(cell_text):
                        ended = True
                        break
                    for para in cell.paragraphs:
                        results.extend(save_from_runs_if_in_range(para.runs))

    return results


def build_vlm_messages_from_docx(docx_path,
                                 prompt_text=(
                                     "请基于给定影像图进行核查并严格按以下要求作答：\n"
                                     "1) 是否标注清楚影像图中的管道位置（实线）及潜在影响半径（虚线）两类线条；\n"
                                     "2) 是否标注管道周边各类人口建筑物和环境受体的名称及相关信息（至少包含人员数量、建筑物名称）。\n"
                                     "请先给出‘是否符合要求：是/否’的结论；如不符合，逐项指出缺失或错误，并给出具体改正建议（应补充哪些线条/标注哪些信息、示例文本）。\n"
                                 ),
                                 start_keyword="一、高后果区基本信息",
                                 end_heading_regex=r"二[、,.．\.]",
                                 output_dir="影像图输出",
                                 first_only=False):
    """将 DOCX 中指定范围的图片转换为 VLM 对话消息结构。

    返回值格式与 `inference/model.py` 中 `make_messages` 一致：
    [
      {
        "role": "user",
        "content": [
          {"type": "image", "image": "path/to/img1.png"},
          {"type": "image", "image": "path/to/img2.png"},
          {"type": "text",  "text":  "...提示词..."}
        ]
      }
    ]
    """
    # 提取图片路径
    items = extract_images_between_headings(
        docx_path,
        start_keyword=start_keyword,
        end_heading_regex=end_heading_regex,
        output_dir=output_dir,
    ) or []

    if first_only and items:
        items = [items[0]]

    content = []
    for it in items:
        img_path = it.get("image_path")
        if img_path:
            content.append({"type": "image", "image": img_path})

    # 无论是否提取到图片，都附加文本提示
    content.append({"type": "text", "text": prompt_text})

    return [
        {
            "role": "user",
            "content": content,
        }
    ]


def parse_args():
    ap = argparse.ArgumentParser(description="识别一、高后果区基本信息到下一个二级标题前的图片")
    ap.add_argument("--input", required=False, default=None, help="输入 DOCX 文件路径；不提供则优先使用样例，否则自动查找")
    ap.add_argument("--output_dir", default="影像图输出", help="图片输出目录")
    ap.add_argument("--keyword", default="一、高后果区基本信息", help="开始分界关键词")
    ap.add_argument("--end_regex", default=r"二[、,.．\.]", help="结束标题正则，默认匹配‘二、’")
    ap.add_argument("--out_path", default=None, help="将识别结果写入该文本/JSON 文件（可选）")
    ap.add_argument("--first_only", action="store_true", help="仅提取范围内第一张图片")
    return ap.parse_args()


def _is_valid_docx_path(p: Path) -> bool:
    return p.suffix.lower() == ".docx" and not p.name.startswith("~$") and p.exists()


def _auto_discover_docx() -> Path | None:
    cwd = Path.cwd()
    candidates = []
    wm = cwd / "word-master"
    if wm.exists():
        for p in wm.rglob("*.docx"):
            if _is_valid_docx_path(p):
                candidates.append(p)
    for p in cwd.glob("*.docx"):
        if _is_valid_docx_path(p):
            candidates.append(p)
    if not candidates:
        for p in cwd.rglob("*.docx"):
            if _is_valid_docx_path(p):
                candidates.append(p)
                break
    return candidates[0] if candidates else None


def main():
    args = parse_args()
    chosen = None
    used_sample = False
    if args.input:
        chosen = Path(args.input)
    else:
        if SAMPLE_FILE_PATH and _is_valid_docx_path(Path(SAMPLE_FILE_PATH)):
            chosen = Path(SAMPLE_FILE_PATH)
            used_sample = True
        else:
            chosen = _auto_discover_docx()
    if not chosen or not Path(chosen).exists():
        print("未提供 --input，且未在当前目录或样例路径找到有效 DOCX（自动忽略临时锁文件 ~$.docx）。请使用 --input 指定文件，例如：")
        print("python yingxiangtu.py --input C:/path/to/file.docx")
        return

    # 统一输出为与 model.make_messages 兼容的 VLM 对话结构
    res_messages = build_vlm_messages_from_docx(
        str(chosen),
        start_keyword=args.keyword,
        end_heading_regex=args.end_regex,
        output_dir=args.output_dir,
        first_only=args.first_only,
    )

    print(json.dumps(res_messages, ensure_ascii=False, indent=2))
    out_file = args.out_path if args.out_path else (SAMPLE_OUT_PATH if used_sample else None)
    if out_file:
        try:
            Path(out_file).parent.mkdir(parents=True, exist_ok=True)
            with open(out_file, "w", encoding="utf-8") as f:
                f.write(json.dumps(res_messages, ensure_ascii=False, indent=2))
        except Exception as e:
            print(f"写入结果失败: {e}")


if __name__ == "__main__":
    main()
