import os
import re
import json
import argparse
from pathlib import Path

# 样例文件与默认输出
SAMPLE_FILE_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/豫洛阳-兰郑长干线-CPY-0790-BFGDGS-ZZSYQFGS.docx"
SAMPLE_OUT_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/image_text_consistency_extract.json"

try:
    from docx import Document
except ImportError:
    raise ImportError("请先 pip install python-docx")


def parse_args():
    ap = argparse.ArgumentParser(description="提取‘一、高后果区基本信息’段的图片与表格中‘高后果区特征描述’文本")
    ap.add_argument("--input", required=False, default=None, help="输入 DOCX 文件路径；不提供则优先使用样例，否则自动查找")
    ap.add_argument("--output_dir", default="影像图输出", help="图片输出目录")
    ap.add_argument("--first_only", action="store_true", help="仅提取范围内第一张图片")
    ap.add_argument("--out_path", default=None, help="将提取的结果写入该 JSON 文件（可选）")
    ap.add_argument("--start_keyword", default="一、高后果区基本信息", help="开始分界关键词")
    ap.add_argument("--end_regex", default=r"二[、,.．\.]", help="结束标题正则，默认匹配‘二、’")
    return ap.parse_args()


def _is_valid_docx_path(p: Path) -> bool:
    return p.suffix.lower() == ".docx" and not p.name.startswith("~$") and p.exists()


def _auto_discover_docx() -> Path | None:
    cwd = Path.cwd()
    candidates = []
    wm = cwd / "word-master"
    if wm.exists():
        for p in wm.rglob("*.docx"):
            if _is_valid_docx_path(p):
                candidates.append(p)
    for p in cwd.glob("*.docx"):
        if _is_valid_docx_path(p):
            candidates.append(p)
    if not candidates:
        for p in cwd.rglob("*.docx"):
            if _is_valid_docx_path(p):
                candidates.append(p)
                break
    return candidates[0] if candidates else None


def extract_feature_text_from_tables(doc_path: str) -> str | None:
    """
    在所有表格中查找包含“高后果区特征描述”的行，返回该行拼接后的文本。
    若未找到则返回 None。
    """
    try:
        doc = Document(doc_path)
    except Exception:
        return None

    headers = {"高后果区基本信息", "高后果区特征描述"}
    for table in doc.tables:
        for row in table.rows:
            row_texts = [c.text.strip() for c in row.cells]
            # 是否存在标识列
            if any("高后果区特征描述" in (t or "") for t in row_texts):
                # 找到“高后果区特征描述”所在的首个单元格索引
                try:
                    hdr_idx = next(i for i, t in enumerate(row_texts) if "高后果区特征描述" in (t or ""))
                except StopIteration:
                    hdr_idx = -1

                # 优先取其右侧的非空、且不是表头词的文本；若无，则取整行中最长的非表头文本
                candidates = []
                for i, t in enumerate(row_texts):
                    if not t:
                        continue
                    # 跳过纯表头词
                    if t in headers:
                        continue
                    # 右侧优先
                    weight = 2 if i > hdr_idx else 1
                    candidates.append((weight, len(t), t))

                if candidates:
                    # 先按是否在右侧、再按长度挑最大者
                    candidates.sort(key=lambda x: (x[0], x[1]), reverse=True)
                    best = candidates[0][2].strip()
                    return best if best else None
    return None


def extract_images_between_headings(doc_path: str,
                                    start_keyword: str,
                                    end_heading_regex: str,
                                    output_dir: str) -> list[dict]:
    """
    提取位于 start_keyword 之后且在下一个符合 end_heading_regex 标题前的所有图片。
    返回: [{"image_path": str, "index": int}]。
    """
    try:
        doc = Document(doc_path)
    except Exception:
        return []

    output_dir_p = Path(output_dir)
    output_dir_p.mkdir(parents=True, exist_ok=True)
    doc_name = Path(doc_path).stem
    in_range = False
    ended = False
    img_idx = 0
    results: list[dict] = []
    end_prog = re.compile(end_heading_regex)

    def save_from_runs_if_in_range(runs):
        nonlocal img_idx
        local = []
        for run in runs:
            drawing = run._element.xpath('.//a:blip')
            for blip in drawing:
                rEmbed = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rEmbed and rEmbed in doc.part.rels:
                    rel = doc.part.rels[rEmbed]
                    if "image" in rel.target_ref:
                        img_idx += 1
                        image_data = rel.target_part.blob
                        img_filename = f"{doc_name}_影像图_{img_idx}.png"
                        img_path = output_dir_p / img_filename
                        with open(img_path, 'wb') as f:
                            f.write(image_data)
                        local.append({"image_path": str(img_path), "index": img_idx})
        return local

    # 扫描正文
    for para in Document(doc_path).paragraphs:
        if ended:
            break
        text = para.text or ""
        if not in_range:
            if start_keyword in text:
                in_range = True
            continue
        if end_prog.search(text):
            ended = True
            break
        results.extend(save_from_runs_if_in_range(para.runs))

    # 扫描表格
    if not ended:
        doc2 = Document(doc_path)
        for table in doc2.tables:
            if ended:
                break
            for row in table.rows:
                if ended:
                    break
                for cell in row.cells:
                    if ended:
                        break
                    cell_text = cell.text or ""
                    if not in_range:
                        if start_keyword in cell_text:
                            in_range = True
                        continue
                    if end_prog.search(cell_text):
                        ended = True
                        break
                    for para in cell.paragraphs:
                        results.extend(save_from_runs_if_in_range(para.runs))

    return results


def build_vlm_messages_from_docx(docx_path: str,
                                 start_keyword: str = "一、高后果区基本信息",
                                 end_heading_regex: str = r"二[、,.．\.]",
                                 output_dir: str = "影像图输出",
                                 first_only: bool = False):
    """
    生成用于图-文一致性核查的 VLM 消息：
    - 图片来源：从“start_keyword”到下一个“二、”标题之前的影像图
    - 文本来源：表格中的“高后果区特征描述”行

    返回：
    - 无图片时：返回纯文本字符串 “未提取到影像图，无需审核。”
    - 有图片时：返回 [ {"role":"user","content":[ {image...}..., {text...} ]} ]
    """
    docx_path = str(docx_path)
    feature_text = extract_feature_text_from_tables(docx_path)
    images = extract_images_between_headings(
        docx_path,
        start_keyword=start_keyword,
        end_heading_regex=end_heading_regex,
        output_dir=output_dir,
    ) or []

    if first_only and images:
        images = images[:1]

    if not images:
        return "未提取到影像图，无需审核。"

    content = []
    for it in images:
        img_path = it.get("image_path")
        if img_path:
            content.append({"type": "image", "image": img_path})

    prompt = (
        "请核查影像图与表格描述的一致性：\n"
        "1) 影像图中是否体现表格‘高后果区特征描述’中的关键要素（位置/对象/数量/范围等）；\n"
        "2) 若不一致，请指出缺失或矛盾之处，并给出需要补充或修正的建议。\n"
        f"表格描述摘录：{feature_text or '（未提取到描述）'}"
    )
    content.append({"type": "text", "text": prompt})

    return [
        {
            "role": "user",
            "content": content,
        }
    ]

def main():
    args = parse_args()
    chosen = None
    used_sample = False
    if args.input:
        chosen = Path(args.input)
    else:
        if SAMPLE_FILE_PATH and _is_valid_docx_path(Path(SAMPLE_FILE_PATH)):
            chosen = Path(SAMPLE_FILE_PATH)
            used_sample = True
        else:
            chosen = _auto_discover_docx()
    if not chosen or not Path(chosen).exists():
        print("未提供 --input，且未在当前目录或样例路径找到有效 DOCX（自动忽略临时文件 ~$.docx）。请用 --input 指定文件，例如：")
        print("python tupian_biaozhu_yizhixing.py --input C:/path/to/file.docx")
        return

    # 1) 提取描述文本（表格‘高后果区特征描述’行）
    feature_text = extract_feature_text_from_tables(str(chosen))

    # 2) 提取图片（默认：从“一、高后果区基本信息”到下一个“二、”标题之间）
    images = extract_images_between_headings(
        str(chosen),
        start_keyword=args.start_keyword,
        end_heading_regex=args.end_regex,
        output_dir=args.output_dir,
    )

    if args.first_only:
        image_paths = [images[0]["image_path"]] if images else []
    else:
        image_paths = [it["image_path"] for it in images]

    result = {
        "docx_path": str(chosen),
        "feature_text": feature_text,
        "image_paths": image_paths,
    }

    print(json.dumps(result, ensure_ascii=False, indent=2))
    out_file = args.out_path if args.out_path else (SAMPLE_OUT_PATH if used_sample else None)
    if out_file:
        try:
            Path(out_file).parent.mkdir(parents=True, exist_ok=True)
            with open(out_file, "w", encoding="utf-8") as f:
                f.write(json.dumps(result, ensure_ascii=False, indent=2))
        except Exception as e:
            print(f"写入结果失败: {e}")


if __name__ == "__main__":
    main()


