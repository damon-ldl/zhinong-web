# -*- coding: utf-8 -*-



import os, re, json
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.table import Table

# ========== 路径配置（你可以改这两个） ==========
BASE_DIR   = r"C:/Users/cassi/Desktop/bisai/word-master"  # 待批量扫描的文件夹
OUT_DIR    = os.path.join(BASE_DIR, "content_check_results")  # 结果输出文件夹

# 如果不存在就建
os.makedirs(OUT_DIR, exist_ok=True)


# ========== 基础工具函数 ==========
def _clean(s: str) -> str:
    return (s or "").strip()

def _norm_value(s: str) -> str:
    """
    归一化：空白去掉、常见分隔统一为'、'、句末标点去掉、'类'=> '型'
    """
    s = (s or "").strip()
    s = s.replace("类", "型")
    s = re.sub(r"[ \t]+", "", s)
    s = re.sub(r"[，,;/|]+", "、", s)
    s = re.sub(r"[：:；;。]$", "", s)
    return s

def _split_types(s: str):
    s = _norm_value(s)
    if not s:
        return []
    parts = [p for p in re.split(r"[、/|，,；; ]+", s) if p]
    seen, out = set(), []
    for p in parts:
        if p not in seen:
            seen.add(p)
            out.append(p)
    return out


# ========== 遍历段落/表格并检测图片 ==========
def iter_block_items(doc):
    """
    按正文顺序遍历文档：段落 -> 表格内段落
    统一返回：
    kind: "paragraph" or "table"
    text: 段落文字
    has_image: 是否检测到图片
    """
    body = doc._element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            # 普通段落
            yield ("paragraph", Paragraph(child, doc), None, None)
        elif isinstance(child, CT_Tbl):
            # 表格 -> 逐格逐段落
            tbl = Table(child, doc)
            for r_idx, row in enumerate(tbl.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p in cell.paragraphs:
                        yield ("table", p, None, (r_idx, c_idx))

def para_has_image(p: Paragraph) -> bool:
    """
    使用 local-name() 规避命名空间问题；检测常见图片相关节点：
    drawing / blip / imagedata / pic / pict / shape
    """
    try:
        if p._element.xpath(
            './/*[local-name()="drawing" or local-name()="blip" or '
            'local-name()="imagedata" or local-name()="pic" or '
            'local-name()="pict" or local-name()="shape"]'
        ):
            return True
        for run in p.runs:
            el = run._element
            if el.xpath(
                './/*[local-name()="drawing" or local-name()="blip" or '
                'local-name()="imagedata" or local-name()="pic" or '
                'local-name()="pict" or local-name()="shape"]'
            ):
                return True
    except Exception:
        pass
    return False


def extract_hca_type_from_tables(doc):
    """
    在表格中查找"高后果区类型"，优先取同一行右侧单元格；
    若为空或极短，再向下取同列若干行（应对换行/合并单元格）。
    """
    for tbl in doc.tables:
        rows = list(getattr(tbl, "rows", []))
        for r_idx, row in enumerate(rows):
            cells = list(getattr(row, "cells", []))
            for c_idx, cell in enumerate(cells):
                raw = _clean(cell.text)
                if "高后果区类型" in raw.replace(" ", ""):
                    # 情况 A：同格内就带值
                    m = re.search(r"高后果区类型[:：]?\s*(.+)", raw)
                    if m and _clean(m.group(1)):
                        return _norm_value(m.group(1))

                    # 情况 B：同一行右侧（去重）
                    right_texts = [_clean(c.text) for c in cells[c_idx + 1:] if _clean(c.text)]
                    # 去重但保持顺序
                    seen, unique_texts = set(), []
                    for txt in right_texts:
                        if txt not in seen:
                            seen.add(txt)
                            unique_texts.append(txt)
                    right_val = "、".join(unique_texts).strip()

                    # 情况 C：同列下一行（取 1~3 行做拼接，兼容换行/合并单元格）
                    down_val = ""
                    if len(right_val) <= 1:
                        frags = []
                        for k in range(1, 4):
                            rr = r_idx + k
                            if rr < len(rows):
                                c2 = rows[rr].cells
                                if c_idx < len(c2):
                                    frags.append(_clean(c2[c_idx].text))
                        # 去重但保持顺序
                        seen, unique_frags = set(), []
                        for f in frags:
                            if f and f not in seen:
                                seen.add(f)
                                unique_frags.append(f)
                        down_val = "、".join(unique_frags).strip()

                    val = right_val if len(right_val) > len(down_val) else down_val
                    val = _norm_value(val)
                    if val:
                        return val
    return None

def load_blocks(doc):
    """
    将 doc 遍历为顺序化 blocks:
    [
      {
        "block_idx": int,
        "kind": "paragraph"/"table",
        "text": "...",
        "has_image": True/False,
        "table_rc": (r,c) or None
      },
      ...
    ]
    """
    blocks = []
    idx = 0
    for kind, p, _, rc in iter_block_items(doc):
        text = _clean(p.text)
        has_img = para_has_image(p)
        blocks.append({
            "block_idx": idx,
            "kind": kind,
            "text": text,
            "has_image": has_img,
            "table_rc": rc
        })
        idx += 1
    return blocks


# ========== 业务规则 ==========
SUPPORTED_TYPES = {"人员密集型", "环境敏感型"}

REQ_MAP = {
    "人员密集型": [
        "高后果区影像图",
        "高后果区现场图",
        "入场线路图",
        "逃生路线图",
        "应急疏散集合点位置",
    ],
    "环境敏感型": [
        "高后果区影像图",
        "高后果区现场图",
        "入场线路图",
        "水体敏感型高后果区围油设施放置图",
    ],
}

ALIASES = {
    "高后果区影像图": [
        "高后果区影像图",
        "影像图",
    ],
    "高后果区现场图": [
        "高后果区现场图",
        "高后果区现场图片",
        "现场图片",
        "现场图",
    ],
    "入场线路图": [
        "入场线路图",
        "入场线路",
    ],
    "逃生路线图": [
        "逃生路线图",
        "逃生路线",
    ],
    "应急疏散集合点位置": [
        "应急疏散集合点位置",
        "应急疏散集结点位置",
        "应急疏散集结点",
        "应急疏散集合点",
        "疏散集结点",
    ],
    "水体敏感型高后果区围油设施放置图": [
        "水体敏感型高后果区围油设施放置图",
        "围油设施放置图",
        "河流流向及围油栏预设点示意图",
        "围油栏预设点示意图",
        "围油栏预设点",
    ],
}

WINDOW = 2  # "标题/说明行"与"图片"最近块的最大距离（段落数）

# === 章节标题匹配（新增） ===
# 通用大标题，如：一、二、三．（兼容中/全角点号与括号序号）
BIG_TITLE_RE = re.compile(r'^\s*[（(]?[一二三四五六七八九十]+[)）]?[、.．]')

# "一、高后果区基本信息"标题（允许前缀编号+多种点号/空格）
BASIC_INFO_TITLE_RE = re.compile(
    r'^\s*[（(]?[一二三四五六七八九十]+[)）]?[、.．]\s*高后果区基本信息\s*$'
)

def near(a, b, w=WINDOW):
    return abs(a - b) <= w

# 编号项检测模式：识别（1）、(2)、1)等格式
NUM_ITEM_PAT = re.compile(r'^\s*[（(]?\s*\d+\s*[)）]')

def is_numbered_item(text: str) -> bool:
    """判断是否是编号项，如（1）、(2)、1)等"""
    return bool(NUM_ITEM_PAT.match(text or ""))


def find_section_range(blocks, title_re=BASIC_INFO_TITLE_RE):
    """
    返回 (start_idx, end_idx, heading_text)
    start_idx: 章节标题所在 block_idx
    end_idx  : 下一章标题的 block_idx（不含），若未找到则为 len(blocks)
    """
    start = None
    for b in blocks:
        txt = (b["text"] or "").strip()
        if title_re.match(txt):
            start = b["block_idx"]
            heading_text = txt
            break
    if start is None:
        return None, None, None

    end = None
    for b in blocks:
        if b["block_idx"] <= start:
            continue
        txt = (b["text"] or "").strip()
        if BIG_TITLE_RE.match(txt):
            end = b["block_idx"]
            break
    if end is None:
        end = (blocks[-1]["block_idx"] + 1) if blocks else 0
    return start, end, heading_text


def find_basic_info_image_evidence(blocks):
    """
    在《一、高后果区基本信息》章节内查找任意图片。
    命中返回 (cap_idx, cap_text, img_idx)，否则返回 None。
    """
    s, e, heading = find_section_range(blocks, BASIC_INFO_TITLE_RE)
    if s is None:
        return None
    # 取该章节内第一张图片
    for b in blocks:
        bi = b["block_idx"]
        if s < bi < e and b["has_image"]:
            cap_idx = s
            cap_text = f"{heading}（检测到章节内图片）"
            img_idx = bi
            return cap_idx, cap_text, img_idx
    return None


# ========== 针对单个文档的检查逻辑（原 main 拆出来） ==========
def analyze_single_doc(doc_path, out_txt_path):
    """
    对单个 docx 执行完整校验，并将结果写入 out_txt_path
    返回最终输出文本。
    """
    if not os.path.exists(doc_path):
        result = f"❌ 找不到文档：{doc_path}"
        with open(out_txt_path, "w", encoding="utf-8") as f:
            f.write(result)
        return result

    # 载入文档
    doc = Document(doc_path)
    blocks = load_blocks(doc)
    # 只用于备用/调试（当前逻辑没直接用 full_text，但保留）
    # full_text = "\n".join(b["text"] for b in blocks)

    # 识别“高后果区类型”（优先基于表格）
    hca_type = extract_hca_type_from_tables(doc)
    all_types = _split_types(hca_type) if hca_type else []
    types_for_check = [t for t in all_types if t in SUPPORTED_TYPES]

    # 先准备输出行
    out_lines = []
    out_lines.append(f"高后果区类型：{hca_type or '未识别'}")

    # 类型不需判断
    if not types_for_check:
        out_lines.append(
            "内容完整性：✅ 无需判断（非人员密集型/环境敏感型或未识别，按规则视为没问题）"
        )
        text = "\n".join(out_lines)
        with open(out_txt_path, "w", encoding="utf-8") as f:
            f.write(text)
        return text

    # 需要严格校验
    # 1) 汇总所有需具备的必需图件（并集去重）
    required = []
    seen_req = set()
    for t in types_for_check:
        for item in REQ_MAP.get(t, []):
            if item not in seen_req:
                seen_req.add(item)
                required.append(item)

    # 2) 找图片 block
    image_idxs = [b["block_idx"] for b in blocks if b["has_image"]]

    # 3) 查找命中的标题/说明 + 就近图片
    evidence_pairs = []  # [(target, cap_idx, cap_text, img_idx)]
    found_targets = set()  # 记录找到的必需图件
    
    for target in required:
        # === 专门规则：高后果区影像图，只要《一、高后果区基本信息》章节下有图片即可 ===
        if target == "高后果区影像图":
            ev = find_basic_info_image_evidence(blocks)
            if ev:
                cap_idx, cap_text, img_idx = ev
                evidence_pairs.append((target, cap_idx, cap_text, img_idx))
                found_targets.add(target)
                continue  # 直接完成该 target 的判定

        # === 其余目标仍使用原有"别名命中 + 向后 WINDOW 范围找图 + 编号项阻断"的严格规则 ===
        aliases = [target] + ALIASES.get(target, [])
        best = None
        for caption in blocks:
            cap_text = caption["text"]
            if not cap_text:
                continue
            # 标题/说明命中别名？
            if not any(alias in cap_text for alias in aliases):
                continue

            cap_idx = caption["block_idx"]

            # ★ 只向"后面"找图片，且限定在 +WINDOW 内
            forward_imgs = [
                ib for ib in image_idxs
                if 0 < ib - cap_idx <= WINDOW
            ]
            if not forward_imgs:
                continue

            # 找到最近的一张候选图
            nb = min(forward_imgs, key=lambda x: x - cap_idx)

            # ★ 若标题与该图片之间出现了"下一条编号项"，说明这张图归下一条，不归当前标题
            has_barrier = any(
                is_numbered_item(blocks[j]["text"])
                for j in range(cap_idx + 1, nb)
                if j < len(blocks)
            )
            if has_barrier:
                continue

            best = (target, cap_idx, cap_text, nb)
            break

        if best:
            evidence_pairs.append(best)
            found_targets.add(target)

    # 4) 检查是否所有必需图件都找到了
    missing_targets = [target for target in required if target not in found_targets]
    
    if missing_targets:
        # 有缺失的图件，报告问题
        out_lines.append("内容完整性：❌ 有问题")
        for missing in missing_targets:
            out_lines.append(f"缺少{missing}")
        out_lines.append("")
        # 仍然输出找到的图件（如果有的话）
        if evidence_pairs:
            out_lines.append("已找到的图件：")
            for target, cap_idx, cap_text, img_idx in evidence_pairs:
                out_lines.append(f"\n【{target}】")
                out_lines.append(cap_text if cap_text else "（无标题文本）")
                out_lines.append("image")
    else:
        # 所有必需图件都找到了
        out_lines.append("内容完整性：✅ 没问题")
        for target, cap_idx, cap_text, img_idx in evidence_pairs:
            out_lines.append(f"\n【{target}】")
            out_lines.append(cap_text if cap_text else "（无标题文本）")
            out_lines.append("image")

    text = "\n".join(out_lines)
    with open(out_txt_path, "w", encoding="utf-8") as f:
        f.write(text)

    return text


# ========== 主调度：批量跑整个文件夹 ==========
def main_batch():
    # 找到 BASE_DIR 下所有 .docx（不含 ~ 开头的临时文件）
    docx_files = [
        f for f in os.listdir(BASE_DIR)
        if f.lower().endswith(".docx") and not f.startswith("~$")
    ]

    if not docx_files:
        print("⚠️ 当前目录未发现 .docx 文件")
        return

    summary = []  # 用于终端打印一个总览
    for fname in docx_files:
        in_path  = os.path.join(BASE_DIR, fname)
        out_name = os.path.splitext(fname)[0] + "_content_consistency_check.txt"
        out_path = os.path.join(OUT_DIR, out_name)

        print("====== 处理文档 ======")
        print(in_path)

        try:
            report_text = analyze_single_doc(in_path, out_path)
            print(report_text)
            print("→ 已写入：", out_path)
            summary.append({
                "file": fname,
                "out": out_name,
                "status": "ok"
            })
        except Exception as e:
            err_msg = f"❌ 处理失败: {e}"
            print(err_msg)
            summary.append({
                "file": fname,
                "out": None,
                "status": err_msg
            })

        print()

    print("====== 批处理完成，总览 ======")
    for item in summary:
        print(f"- {item['file']}: {item['status']} -> {item.get('out')}")


if __name__ == "__main__":
    main_batch()

# ========== 供外部调用的无副作用函数 ==========
def build_content_consistency_from_docx(doc_path: str):
    """
    针对单个 .docx 返回内容一致性检查的人类可读文本（不写文件，不调用 VLM）。
    """
    try:
        if not os.path.exists(doc_path):
            msg = f"❌ 找不到文档：{doc_path}"
            return msg

        doc = Document(doc_path)
        blocks = load_blocks(doc)

        hca_type = extract_hca_type_from_tables(doc)
        all_types = _split_types(hca_type) if hca_type else []
        types_for_check = [t for t in all_types if t in SUPPORTED_TYPES]

        out_lines = []
        out_lines.append(f"高后果区类型：{hca_type or '未识别'}")

        # 类型不需判断
        if not types_for_check:
            out_lines.append(
                "内容完整性：✅ 无需判断（非人员密集型/环境敏感型或未识别，按规则视为没问题）"
            )
            text = "\n".join(out_lines)
            return text

        # 严格校验
        required = []
        seen_req = set()
        for t in types_for_check:
            for item in REQ_MAP.get(t, []):
                if item not in seen_req:
                    seen_req.add(item)
                    required.append(item)

        image_idxs = [b["block_idx"] for b in blocks if b["has_image"]]

        evidence_pairs = []  # [(target, cap_idx, cap_text, img_idx)]
        found_targets = set()

        for target in required:
            if target == "高后果区影像图":
                ev = find_basic_info_image_evidence(blocks)
                if ev:
                    cap_idx, cap_text, img_idx = ev
                    evidence_pairs.append((target, cap_idx, cap_text, img_idx))
                    found_targets.add(target)
                    continue

            aliases = [target] + ALIASES.get(target, [])
            best = None
            for caption in blocks:
                cap_text = caption["text"]
                if not cap_text:
                    continue
                if not any(alias in cap_text for alias in aliases):
                    continue
                cap_idx = caption["block_idx"]
                forward_imgs = [
                    ib for ib in image_idxs
                    if 0 < ib - cap_idx <= WINDOW
                ]
                if not forward_imgs:
                    continue
                nb = min(forward_imgs, key=lambda x: x - cap_idx)
                has_barrier = any(
                    is_numbered_item(blocks[j]["text"])
                    for j in range(cap_idx + 1, nb)
                    if j < len(blocks)
                )
                if has_barrier:
                    continue
                best = (target, cap_idx, cap_text, nb)
                break

            if best:
                evidence_pairs.append(best)
                found_targets.add(target)

        missing_targets = [target for target in required if target not in found_targets]

        if missing_targets:
            out_lines.append("内容完整性：❌ 有问题")
            for missing in missing_targets:
                out_lines.append(f"缺少{missing}")
            out_lines.append("")
            if evidence_pairs:
                out_lines.append("已找到的图件：")
                for target, cap_idx, cap_text, img_idx in evidence_pairs:
                    out_lines.append(f"\n【{target}】")
                    out_lines.append(cap_text if cap_text else "（无标题文本）")
                    out_lines.append("image")
            text = "\n".join(out_lines)
            return text
        else:
            out_lines.append("内容完整性：✅ 没问题")
            for target, cap_idx, cap_text, img_idx in evidence_pairs:
                out_lines.append(f"\n【{target}】")
                out_lines.append(cap_text if cap_text else "（无标题文本）")
                out_lines.append("image")
            text = "\n".join(out_lines)
            return text
    except Exception as e:
        return f"❌ 处理失败: {e}"