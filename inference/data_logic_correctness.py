import os
import re
import json
import argparse
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ImportError:
    raise ImportError("请先 pip install python-docx")

# 可选 PDF 支持
try:
    import pdfplumber  # type: ignore
except Exception:
    pdfplumber = None  # 延迟到使用时再提示安装


# 与项目风格一致的样例与默认输出
SAMPLE_FILE_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/豫洛阳-兰郑长干线-CPY-0790-BFGDGS-ZZSYQFGS.docx"
SAMPLE_OUT_PATH_POTENTIAL = r"C:/Users/cassi/Desktop/bisai/word-master/data_logic_correctness_potential.json"
SAMPLE_OUT_PATH_RISK = r"C:/Users/cassi/Desktop/bisai/word-master/data_logic_correctness_risk.json"


def parse_args():
    ap = argparse.ArgumentParser(description="从 DOCX/PDF 提取数据并检查数据逻辑是否正确（电位范围、风险字段数值性、风险等级有效性等）")
    ap.add_argument("--input", required=False, default=None, help="输入 DOCX/PDF 文件路径；不提供则优先使用样例，否则自动查找")
    ap.add_argument("--out_path_potential", default=None, help="电位结果 JSON 文件路径（可选）")
    ap.add_argument("--out_path_risk", default=None, help="风险结果 JSON 文件路径（可选）")
    ap.add_argument("--start_keyword", default=None, help="可选：仅在包含该关键词之后的范围进行粗过滤")
    ap.add_argument("--end_regex", default=None, help="可选：遇到该正则命中的标题/文本即结束范围")
    return ap.parse_args()


def _is_valid_input_path(p: Path) -> bool:
    suf = p.suffix.lower()
    if p.name.startswith("~$"):
        return False
    return suf in {".docx", ".pdf"} and p.exists()


def _auto_discover_input() -> Path | None:
    cwd = Path.cwd()
    candidates: list[Path] = []
    wm = cwd / "word-master"
    patterns = ["*.docx", "*.pdf"]
    if wm.exists():
        for pat in patterns:
            for p in wm.rglob(pat):
                if _is_valid_input_path(p):
                    candidates.append(p)
    for pat in patterns:
        for p in cwd.glob(pat):
            if _is_valid_input_path(p):
                candidates.append(p)
    if not candidates:
        for pat in patterns:
            for p in cwd.rglob(pat):
                if _is_valid_input_path(p):
                    candidates.append(p)
                    break
    return candidates[0] if candidates else None


def extract_text_from_docx(docx_path: str) -> str:
    try:
        doc = Document(docx_path)
    except Exception:
        return ""

    full_text: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            full_text.append(t)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    row_text.append(ct)
            if row_text:
                full_text.append("  ".join(row_text))

    text = "\n".join(full_text)
    return text


def extract_text_from_pdf(pdf_path: str) -> str:
    if pdfplumber is None:
        raise ImportError("检测到 PDF 输入，但未安装 pdfplumber。请先 pip install pdfplumber")
    try:
        full_text: list[str] = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                t = (page.extract_text() or "").strip()
                if t:
                    full_text.append(t)
        text = "\n".join(full_text)
        return text
    except Exception:
        return ""


def extract_text_from_file(path: str) -> str:
    suf = Path(path).suffix.lower()
    if suf == ".docx":
        return extract_text_from_docx(path)
    if suf == ".pdf":
        return extract_text_from_pdf(path)
    return ""


# ========================== 规则与提取 ==========================

VALIDATION_RULES = {
    "potential_range": (-1.2, -0.85),
    "risk_levels": ["低", "中", "较高", "高"],
}


def extract_potential_from_table(doc_path: str) -> list[dict[str, Any]]:
    """
    从"高后果区管道电位测试结果"表格中提取电位值。
    返回元素：{"value": float, "context": str}
    """
    potential_results: list[dict[str, Any]] = []
    
    try:
        doc = Document(doc_path)
        full_content = ""
        
        # 提取段落文本（用于定位表格标题）
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                full_content += para_text + "\n"
        
        # 定位"高后果区管道电位测试结果"标题及对应表格
        table_title = "高后果区管道电位测试结果"
        if table_title not in full_content:
            return potential_results  # 未找到表格，返回空列表
        
        # 找到标题所在段落索引
        title_para_index = None
        for i, para in enumerate(doc.paragraphs):
            if table_title in para.text.strip():
                title_para_index = i
                break
        
        if title_para_index is None:
            return potential_results
        
        # 从标题段落后查找表格
        target_table = None
        for i in range(title_para_index + 1, len(doc.paragraphs)):
            if doc.paragraphs[i].text.strip() == "":
                continue
            # 遍历所有表格，确认与标题关联性
            for table in doc.tables:
                if len(table.rows) == 0:
                    continue
                first_row = "|".join([cell.text.strip() for cell in table.rows[0].cells])
                if "测试桩号" in first_row or "电位" in first_row:
                    target_table = table
                    break
            if target_table:
                break
        
        if not target_table:
            return potential_results
        
        # 提取列名，找到电位列索引
        header = [cell.text.strip().replace("\n", " ") for cell in target_table.rows[0].cells]
        potential_col_indices = []
        for idx, col_name in enumerate(header):
            if "电位" in col_name or "V" in col_name or "v" in col_name:
                potential_col_indices.append(idx)
        
        # 如果没有明确的电位列，尝试从所有数值列中提取
        if not potential_col_indices:
            # 假设除第一列（通常是编号/名称）外都可能是电位值
            potential_col_indices = list(range(1, len(header)))
        
        # 提取数据行中的电位值
        for row_idx, row in enumerate(target_table.rows[1:], start=1):
            row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if not any(cell != "" for cell in row_data):
                continue
            
            # 从电位列中提取数值
            row_potentials = []  # 用于去重
            for col_idx in potential_col_indices:
                if col_idx >= len(row_data):
                    continue
                cell_text = row_data[col_idx].strip()
                if not cell_text:
                    continue
                
                # 尝试提取数值（可能包含多个值，用空格或逗号分隔）
                numbers = re.findall(r"([+\-]?\d+(?:\.\d+)?)", cell_text)
                for num_str in numbers:
                    try:
                        val = float(num_str)
                        # 限制在合理范围内
                        if -2.0 <= val <= 2.0:
                            found_in_cols = True
                            # 去重：避免同一个值被提取多次
                            if val not in row_potentials:
                                row_potentials.append(val)
                                # 构建上下文
                                context = f"表格：{table_title} | 列：{header[col_idx] if col_idx < len(header) else ''} | 行{row_idx}：{' | '.join(row_data)}"
                                potential_results.append({"value": val, "context": context})
                    except (ValueError, TypeError):
                        continue
            
            # 如果从列中没有找到，且整行看起来像电位数据（例如：K79 +0.989 +0.994 +1.002）
            if not row_potentials:
                row_text = " ".join(row_data)
                if "K" in row_text or any(c in row_text for c in ["+", "-"]):
                    numbers = re.findall(r"([+\-]?\d+(?:\.\d+)?)", row_text)
                    for num_str in numbers:
                        try:
                            val = float(num_str)
                            if -2.0 <= val <= 2.0:
                                context = f"表格：{table_title} | 行{row_idx}：{row_text}"
                                potential_results.append({"value": val, "context": context})
                        except (ValueError, TypeError):
                            continue
        
    except Exception:
        # 表格提取失败，返回空列表，后续会用文本提取作为备选
        pass
    
    return potential_results


def extract_potential_test_results(text: str, doc_path: str | None = None) -> list[dict[str, Any]]:
    """
    提取电位测试结果（单位V或含正负号，范围在[-2, 2]之间）。
    优先从"高后果区管道电位测试结果"表格中提取（如果是DOCX文件），
    否则从文本中提取。
    兼容示例：
    - 电位：-0.95V
    - K79 +0.989 +0.994 +1.002
    - 近6月 0.97 0.98 1.00
    返回元素：{"value": float, "context": str}
    """
    potential_results: list[dict[str, Any]] = []
    
    # 优先尝试从表格提取（仅对DOCX文件）
    if doc_path and Path(doc_path).suffix.lower() == ".docx":
        table_results = extract_potential_from_table(doc_path)
        if table_results:
            potential_results.extend(table_results)
    
    # 如果表格提取为空或不是DOCX，使用文本提取作为备选
    if not potential_results:
        clean_text = re.sub(r"\s+", " ", text or "")
        
        patterns = [
            r"(?:电位|IR降|K\d+)[^0-9\-+]{0,10}?([+\-]?\d+(?:\.\d+)?)\s*[Vv伏]?",
            r"(?:电位测试结果|近\d+月)[：:\s]*([+\-]?\d+(?:\.\d+)?)",
        ]
        
        for pattern in patterns:
            for m in re.finditer(pattern, clean_text):
                try:
                    val = float(m.group(1))
                    if not (-2.0 <= val <= 2.0):
                        continue
                    context = clean_text[max(0, m.start() - 80) : m.end() + 80]
                    potential_results.append({"value": val, "context": context.strip()})
                except Exception:
                    continue
    
    return potential_results


def extract_risk_assessment_data(text: str) -> dict:
    """
    提取风险评价数据：可能性、后果值、风险值（数值）与风险等级（枚举）。
    返回：
    {
      "possibility_values": [{"value": float, "raw_text": str, "context": str}],
      "consequence_values": [...],
      "risk_values": [...],
      "risk_levels": [{"level": str, "raw_text": str, "context": str}],
    }
    """
    clean_text = re.sub(r"\s+", " ", text or "")
    risk_data = {
        "possibility_values": [],
        "consequence_values": [],
        "risk_values": [],
        "risk_levels": [],
    }

    for match in re.finditer(r"(?:失效)?可能性[：:\s]*([0-9]*\.?[0-9]+)", clean_text):
        risk_data["possibility_values"].append(
            {
                "value": float(match.group(1)),
                "raw_text": match.group(0),
                "context": clean_text[max(0, match.start() - 80) : match.end() + 80],
            }
        )

    for match in re.finditer(r"(?:失效)?后果值[：:\s]*([0-9]*\.?[0-9]+)", clean_text):
        risk_data["consequence_values"].append(
            {
                "value": float(match.group(1)),
                "raw_text": match.group(0),
                "context": clean_text[max(0, match.start() - 80) : match.end() + 80],
            }
        )

    for match in re.finditer(r"风险值[：:\s]*([0-9]*\.?[0-9]+)", clean_text):
        risk_data["risk_values"].append(
            {
                "value": float(match.group(1)),
                "raw_text": match.group(0),
                "context": clean_text[max(0, match.start() - 80) : match.end() + 80],
            }
        )

    for match in re.finditer(r"风险等级[：:\s]*([低中较高高]+)", clean_text):
        risk_data["risk_levels"].append(
            {
                "level": match.group(1).strip(),
                "raw_text": match.group(0),
                "context": clean_text[max(0, match.start() - 80) : match.end() + 80],
            }
        )

    return risk_data


def verify_data_logic(potential_results: list[dict[str, Any]], risk_data: dict) -> dict:
    """
    验证数据逻辑：
    - 电位值是否在[-1.2, -0.85]V 之间
    - 可能性/后果值/风险值是否为数值
    - 风险等级是否在 {低, 中, 较高, 高}
    - 关键数据缺失提示
    返回包含评分、问题列表与汇总统计
    """
    verification_results = {
        "potential_valid": True,
        "risk_values_numeric": True,
        "risk_levels_valid": True,
        "data_logic_score": 5,
        "issues": [],
        "data_summary": {},
    }

    valid_potentials = 0
    invalid_potentials: list[dict[str, Any]] = []
    for result in potential_results:
        value = result.get("value")
        if value is None:
            continue
        if VALIDATION_RULES["potential_range"][0] <= value <= VALIDATION_RULES["potential_range"][1]:
            valid_potentials += 1
        else:
            invalid_potentials.append(result)

    if invalid_potentials:
        verification_results["potential_valid"] = False
        verification_results["issues"].append(
            f"电位测试结果超出范围：{[r['value'] for r in invalid_potentials]}"
        )
        verification_results["data_logic_score"] -= 2

    verification_results["data_summary"]["potential_results"] = {
        "total": len(potential_results),
        "valid": valid_potentials,
        "invalid": len(invalid_potentials),
    }

    numeric_checks = {
        "possibility": risk_data.get("possibility_values", []),
        "consequence": risk_data.get("consequence_values", []),
        "risk_value": risk_data.get("risk_values", []),
    }
    for field_name, values in numeric_checks.items():
        non_numeric = []
        for item in values:
            if not isinstance(item.get("value"), (int, float)):
                non_numeric.append(item)
        if non_numeric:
            verification_results["risk_values_numeric"] = False
            verification_results["issues"].append(
                f"{field_name}字段包含非数值：{[item.get('raw_text') for item in non_numeric]}"
            )
            verification_results["data_logic_score"] -= 1

    verification_results["data_summary"]["risk_numeric_fields"] = {
        "possibility_count": len(risk_data.get("possibility_values", [])),
        "consequence_count": len(risk_data.get("consequence_values", [])),
        "risk_value_count": len(risk_data.get("risk_values", [])),
    }

    invalid_levels = []
    for level_item in risk_data.get("risk_levels", []):
        level = level_item.get("level")
        if level not in VALIDATION_RULES["risk_levels"]:
            invalid_levels.append(level_item)
    if invalid_levels:
        verification_results["risk_levels_valid"] = False
        verification_results["issues"].append(
            f"风险等级不符合标准：{[item.get('level') for item in invalid_levels]}"
        )
        verification_results["data_logic_score"] -= 1

    verification_results["data_summary"]["risk_levels"] = {
        "total": len(risk_data.get("risk_levels", [])),
        "valid": len(risk_data.get("risk_levels", [])) - len(invalid_levels),
        "invalid": len(invalid_levels),
    }

    if not potential_results:
        verification_results["issues"].append("缺少电位测试结果数据")
        verification_results["data_logic_score"] -= 1
    if not any(
        [
            risk_data.get("possibility_values"),
            risk_data.get("consequence_values"),
            risk_data.get("risk_values"),
        ]
    ):
        verification_results["issues"].append("缺少风险评价数值数据")
        verification_results["data_logic_score"] -= 1
    if not risk_data.get("risk_levels"):
        verification_results["issues"].append("缺少风险等级数据")
        verification_results["data_logic_score"] -= 1

    verification_results["data_logic_score"] = max(0, verification_results["data_logic_score"])  # 保底不为负
    return verification_results


def extract_risk_table_content(doc_path: str) -> str:
    """
    参考 luoji.py 的逻辑，提取"高后果区风险评价结果"表格的原始内容。
    返回表格文本内容，如果未找到则返回提示信息。
    """
    try:
        doc = Document(doc_path)
        keywords = ['失效可能性', '失效后果值', '风险值', '风险等级']
        result_rows = []

        for table in doc.tables:
            for row in table.rows:
                row_data = []
                seen = set()  # 用于跟踪已出现的内容
                for cell in row.cells:
                    # 清理单元格文本（去除多余空白）
                    clean_text = re.sub(r'\s+', ' ', cell.text.strip())
                    # 只保留未出现过的内容，保持原始顺序
                    if clean_text not in seen:
                        seen.add(clean_text)
                        row_data.append(clean_text)

                # 检查是否包含关键词且避免重复行
                has_keyword = any(kw in row_data for kw in keywords)
                if has_keyword and row_data not in result_rows:
                    result_rows.append(row_data)

        if not result_rows:
            return f"文件\"{Path(doc_path).name}\"中未找到包含风险评价关键词的表格"

        # 构造输出内容
        lines = []
        lines.append("表格名：高后果区风险评价结果")
        lines.append("表格内容：")
        for idx, row in enumerate(result_rows, 1):
            lines.append(f"第{idx}行：" + " | ".join(row))
        
        return "\n".join(lines)
        
    except Exception as e:
        return f"提取风险表格时出错：{str(e)}"


def extract_potential_table_content(doc_path: str) -> str:
    """
    参考 dianwei.py 的逻辑，提取"高后果区管道电位测试结果"表格的原始内容。
    返回表格文本内容，如果未找到则返回提示信息。
    """
    try:
        doc = Document(doc_path)
        full_content = ""
        
        # 提取段落文本（用于定位表格标题）
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                full_content += para_text + "\n"
        
        # 定位"高后果区管道电位测试结果"标题及对应表格
        table_title = "高后果区管道电位测试结果"
        if table_title not in full_content:
            return f"文件\"{Path(doc_path).name}\"中未找到\"{table_title}\"表"
        
        # 找到标题所在段落索引
        title_para_index = None
        for i, para in enumerate(doc.paragraphs):
            if table_title in para.text.strip():
                title_para_index = i
                break
        
        if title_para_index is None:
            return f"文件\"{Path(doc_path).name}\"中未找到\"{table_title}\"标题"
        
        # 从标题段落后查找表格
        target_table = None
        for i in range(title_para_index + 1, len(doc.paragraphs)):
            if doc.paragraphs[i].text.strip() == "":
                continue
            # 遍历所有表格，确认与标题关联性
            for table in doc.tables:
                if len(table.rows) == 0:
                    continue
                first_row = "|".join([cell.text.strip() for cell in table.rows[0].cells])
                if "测试桩号" in first_row or "电位" in first_row:
                    target_table = table
                    break
            if target_table:
                break
        
        if not target_table:
            return f"文件\"{Path(doc_path).name}\"中\"{table_title}\"标题后未找到对应表格"
        
        # 提取表格所有内容（列名+数据）
        header = [cell.text.strip().replace("\n", " ") for cell in target_table.rows[0].cells]
        data = []
        for row in target_table.rows[1:]:
            row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cell != "" for cell in row_data):
                data.append(row_data)
        
        # 构造输出内容
        lines = []
        lines.append(f"表格名：{table_title}")
        lines.append("表格列名：" + " | ".join(header))
        lines.append("表格数据：")
        for idx, row in enumerate(data, 1):
            lines.append(f"第{idx}行：" + " | ".join(row))
        
        return "\n".join(lines)
        
    except Exception as e:
        return f"提取表格时出错：{str(e)}"


def build_potential_json(doc_path: str) -> dict:
    """
    构建电位测试结果的 JSON 输出，参考 luxian_extract.py 格式。
    返回包含 messages 的字典。
    """
    p = Path(doc_path)
    if not p.exists() or not p.is_file():
        return {
            "messages": [
                {
                    "content": f"❌ 找不到文档：{doc_path}",
                    "role": "user"
                }
            ]
        }

    # 提取表格内容（仅对 DOCX 文件）
    if p.suffix.lower() == ".docx":
        table_content = extract_potential_table_content(str(p))
        content = f"内容：\n{table_content}\n\n帮我审核电位测试结果部分是否超出-0.85V~-1.2V"
    else:
        # PDF 或其他格式，回退到文本提取
        text = extract_text_from_file(str(p))
        potential_results = extract_potential_test_results(text, doc_path=str(p))
        
        content_parts = ["内容："]
        if potential_results:
            content_parts.append("电位测试结果：")
            for item in potential_results:
                v = item.get("value")
                ctx = item.get("context", "")
                content_parts.append(f"{v} V | {ctx}")
        else:
            content_parts.append("未找到电位测试数据")
        content_parts.append("")
        content_parts.append("帮我审核电位测试结果部分是否超出-0.85V~-1.2V")
        content = "\n".join(content_parts)

    return {
        "messages": [
            {
                "content": content,
                "role": "user"
            }
        ]
    }


def build_risk_json(doc_path: str) -> dict:
    """
    构建风险评价结果的 JSON 输出，参考 luoji.py 和 luxian_extract.py 格式。
    返回包含 messages 的字典。
    """
    p = Path(doc_path)
    if not p.exists() or not p.is_file():
        return {
            "messages": [
                {
                    "content": f"❌ 找不到文档：{doc_path}",
                    "role": "user"
                }
            ],
            "images": []
        }

    # 提取表格内容（仅对 DOCX 文件）
    if p.suffix.lower() == ".docx":
        table_content = extract_risk_table_content(str(p))
        # 提取表格内容部分（保留完整表格信息，但确保格式正确）
        if "表格名：" in table_content and "表格内容：" in table_content:
            # 已经包含表格名和表格内容标签，直接使用
            content = f"{table_content}\n帮我审核风险评价结果中可能性、后果值、风险值数据是否为数值、风险等级是否为'低、中、较高、高'。"
        else:
            # 如果没有找到标准格式，添加表格名和标签
            content = f"表格名：高后果区风险评价结果\n表格内容：{table_content}\n帮我审核风险评价结果中可能性、后果值、风险值数据是否为数值、风险等级是否为'低、中、较高、高'。"
    else:
        # PDF 或其他格式，回退到文本提取
        text = extract_text_from_file(str(p))
        risk_data = extract_risk_assessment_data(text)
        
        content_parts = ["表格名：高后果区风险评价结果", "表格内容："]
        if risk_data.get("possibility_values") or risk_data.get("consequence_values") or risk_data.get("risk_values") or risk_data.get("risk_levels"):
            for item in risk_data.get("possibility_values", []):
                content_parts.append(f"失效可能性：{item.get('value')}")
            for item in risk_data.get("consequence_values", []):
                content_parts.append(f"失效后果值：{item.get('value')}")
            for item in risk_data.get("risk_values", []):
                content_parts.append(f"风险值：{item.get('value')}")
            for item in risk_data.get("risk_levels", []):
                content_parts.append(f"风险等级：{item.get('level')}")
        else:
            content_parts.append("未找到风险评价数据")
        content_parts.append("帮我审核风险评价结果中可能性、后果值、风险值数据是否为数值、风险等级是否为'低、中、较高、高'。")
        content = "\n".join(content_parts)

    return {
        "messages": [
            {
                "content": content,
                "role": "user"
            }
        ]
    }


def build_data_logic_correctness_from_file(
    doc_path: str,
    start_keyword: str | None = None,
    end_regex: str | None = None,
) -> str:
    p = Path(doc_path)
    if not p.exists() or not p.is_file():
        return f"❌ 找不到文档：{doc_path}"

    # 提取表格内容（仅对 DOCX 文件）
    if p.suffix.lower() == ".docx":
        table_content = extract_potential_table_content(str(p))
        lines = []
        lines.append("内容：")
        lines.append(table_content)
        lines.append("")
        lines.append("帮我审核电位测试结果部分是否超出-0.85V~-1.2V")
        return "\n".join(lines)
    else:
        # PDF 或其他格式，回退到文本提取
        text = extract_text_from_file(str(p))
        if start_keyword:
            idx = text.find(start_keyword)
            if idx >= 0:
                text = text[idx:]
        if end_regex:
            m = re.search(end_regex, text)
            if m:
                text = text[: m.start()]
        
        potential_results = extract_potential_test_results(text, doc_path=str(p))
        
        lines = []
        lines.append("内容：")
        if potential_results:
            lines.append("电位测试结果：")
            for item in potential_results:
                v = item.get("value")
                ctx = item.get("context", "")
                lines.append(f"{v} V | {ctx}")
        else:
            lines.append("未找到电位测试数据")
        
        lines.append("")
        lines.append("帮我审核电位测试结果部分是否超出-0.85V~-1.2V")
        
        return "\n".join(lines)

def main():
    args = parse_args()
    chosen = None
    
    if args.input:
        chosen = Path(args.input)
    else:
        if SAMPLE_FILE_PATH and _is_valid_input_path(Path(SAMPLE_FILE_PATH)):
            chosen = Path(SAMPLE_FILE_PATH)
        else:
            chosen = _auto_discover_input()

    if not chosen or not Path(chosen).exists():
        print("未提供 --input，且未在当前目录或样例路径找到有效 DOCX/PDF（自动忽略临时文件 ~$.docx）。")
        print("请使用 --input 指定文件，例如：")
        print("python data_logic_correctness.py --input C:/path/to/file.docx")
        return

    # 生成电位 JSON
    potential_json = build_potential_json(str(chosen))
    print("电位 JSON:")
    print(json.dumps([potential_json], ensure_ascii=False, indent=2))

    # 生成风险 JSON
    risk_json = build_risk_json(str(chosen))
    print("\n风险 JSON:")
    print(json.dumps([risk_json], ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()


