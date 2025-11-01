# -*- coding: utf-8 -*-
"""
PDF 关键词定位裁剪 + 首个签字检测（编制/校对/审核）
DOCX 图片提取 + 首个命中即停
"""

import os
import re
import io
import json
import base64
import zipfile
import argparse
from pathlib import Path
from typing import List, Dict, Tuple, Optional

# ------- 第三方库 -------
# Word
try:
    from docx import Document
except ImportError:
    Document = None

# PDF
import fitz  # PyMuPDF
from PIL import Image

# VLM API
import requests

# 时间戳
try:
    import pandas as pd
except ImportError:
    pd = None
from datetime import datetime


# 样例文件路径（与 yingxiangtu.py 保持一致的测试文档）
SAMPLE_FILE_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/豫洛阳-兰郑长干线-CPY-0790-BFGDGS-ZZSYQFGS.docx"


# ===================== 工具类：DOCX / PDF 基础图片提取 =====================
class DocumentImageExtractor:
    """文档图片提取器（保持你的原逻辑，供 DOCX 与 PDF 回退使用）"""

    def __init__(self, output_dir: str = "extracted_images"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def extract_from_word(self, doc_path: str) -> List[Dict]:
        """
        只提取 Word 文件中“一、高后果区基本信息”前所有图片。
        支持段落、表格、run中drawing图片整体提取。
        """
        from docx.oxml.ns import qn
        images = []
        doc_name = Path(doc_path).stem
        try:
            if Document is None:
                raise ImportError("未安装 python-docx，请先 pip install python-docx")
            doc = Document(doc_path)
            img_count = 0
            found_break = False
            def extract_imgs_from_runs(runs):
                nonlocal img_count
                for run in runs:
                    drawing = run._element.xpath('.//a:blip')
                    for blip in drawing:
                        rEmbed = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rEmbed and rEmbed in doc.part.rels:
                            rel = doc.part.rels[rEmbed]
                            if "image" in rel.target_ref:
                                img_count += 1
                                image_data = rel.target_part.blob
                                img_filename = f"{doc_name}_advanced_img_{img_count}.png"
                                img_path = self.output_dir / img_filename
                                with open(img_path, 'wb') as f:
                                    f.write(image_data)
                                images.append({
                                    'source_file': doc_path,
                                    'image_path': str(img_path),
                                    'image_type': 'embedded',
                                    'index': img_count,
                                    'stop_before': '一、高后果区基本信息'
                                })
            # 遍历所有正文段落
            for para in doc.paragraphs:
                if "一、高后果区基本信息" in para.text:
                    found_break = True
                    break
                extract_imgs_from_runs(para.runs)
            if not found_break:
                # 遍历所有表格
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            # 若出现分界词则停止
                            if "一、高后果区基本信息" in cell.text:
                                found_break = True
                                break
                            for para in cell.paragraphs:
                                extract_imgs_from_runs(para.runs)
                        if found_break:
                            break
                    if found_break:
                        break
        except Exception as e:
            print(f"[extract_from_word] 出错: {e}")
        return images

    def extract_from_pdf(self, pdf_path: str) -> List[Dict]:
        """从 PDF 提取页面内嵌图片（作为回退方案）"""
        images = []
        doc_name = Path(pdf_path).stem
        try:
            pdf_doc = fitz.open(pdf_path)
            img_count = 0
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    img_count += 1
                    xref = img[0]
                    pix = fitz.Pixmap(pdf_doc, xref)
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        img_data = pix.tobytes("png")
                        img_filename = f"{doc_name}_page{page_num+1}_img{img_index+1}.png"
                        img_path = self.output_dir / img_filename
                        with open(img_path, 'wb') as f:
                            f.write(img_data)
                        images.append({
                            'source_file': pdf_path,
                            'image_path': str(img_path),
                            'image_type': 'page_image',
                            'page': page_num + 1,
                            'index': img_count
                        })
                    pix = None
            pdf_doc.close()
        except Exception as e:
            print(f"[extract_from_pdf] 出错: {e}")
        return images

    def extract_images_above_keyword(self, doc_path: str, keyword_regex: str) -> List[Dict]:
        """
        提取 Word 文件中“编制单位”(支持正则)段落前所有图片（正文和表格均处理）。
        """
        import re
        from docx.oxml.ns import qn
        images = []
        doc_name = Path(doc_path).stem
        try:
            if Document is None:
                raise ImportError("未安装 python-docx，请先 pip install python-docx")
            doc = Document(doc_path)
            img_count = 0
            found_break = False
            prog = re.compile(keyword_regex)
            def extract_imgs_from_runs(runs):
                nonlocal img_count
                for run in runs:
                    drawing = run._element.xpath('.//a:blip')
                    for blip in drawing:
                        rEmbed = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rEmbed and rEmbed in doc.part.rels:
                            rel = doc.part.rels[rEmbed]
                            if "image" in rel.target_ref:
                                img_count += 1
                                image_data = rel.target_part.blob
                                img_filename = f"{doc_name}_advanced_img_{img_count}.png"
                                img_path = self.output_dir / img_filename
                                with open(img_path, 'wb') as f:
                                    f.write(image_data)
                                images.append({
                                    'source_file': doc_path,
                                    'image_path': str(img_path),
                                    'image_type': 'embedded',
                                    'index': img_count,
                                    'stop_before': keyword_regex
                                })
            # 正文(不含表格)
            for para in doc.paragraphs:
                if prog.search(para.text):
                    found_break = True
                    break
                extract_imgs_from_runs(para.runs)
            if not found_break:
                # 表格内
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if prog.search(cell.text):
                                found_break = True
                                break
                            for para in cell.paragraphs:
                                extract_imgs_from_runs(para.runs)
                        if found_break:
                            break
                    if found_break:
                        break
        except Exception as e:
            print(f"[extract_images_above_keyword] 出错: {e}")
        return images

    def extract_above_gqgqxx_and_check_bzdw_images(self, doc_path: str, gqgqxx_regex: str, bzdw_regex: str) -> dict:
        """
        1. 定位“一、高后果区基本信息”分界，只分析其上方全部内容
        2. 在其上方内容里，继续找“编制单位”关键字，统计其上方图片数量。
        3. 返回 {'all_images_above_gqgqxx': [...], 'bzdw_images_count': int, 'bzdw_images': [...]}。
        """
        import re
        from docx.oxml.ns import qn
        doc_name = Path(doc_path).stem
        images = []
        bzdw_images = []
        img_count = 0
        bzdw_count = 0
        found_gqgq_idx = None
        found_bzdw_idx = None
        gqgqxx_prog = re.compile(gqgqxx_regex)
        bzdw_prog = re.compile(bzdw_regex)
        # step1-收集所有块（正文和表格cell）作为线性列表
        structure = []  # [(type, text, para/cell_obj)]
        try:
            if Document is None:
                raise ImportError("未安装 python-docx，请先 pip install python-docx")
            doc = Document(doc_path)
            # 正文
            for para in doc.paragraphs:
                structure.append(("para", para.text, para))
            # 表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        structure.append(("cell", cell.text, cell))
            # 找gqgqxx分界
            for idx, (_, text, _) in enumerate(structure):
                if gqgqxx_prog.search(text):
                    found_gqgq_idx = idx
                    break
            if found_gqgq_idx is None:
                # 没找到直接退，默认全范围
                relevant = structure
            else:
                # 只要高后果区基本信息上方
                relevant = structure[:found_gqgq_idx]
            # 在 relevant 范围找“编制单位”分界
            for idx, (_, text, _) in enumerate(relevant):
                if bzdw_prog.search(text):
                    found_bzdw_idx = idx
                    break
            # step2-遍历 relevant, 每遇图片都存，同时顺便判断bzdw前所含图片
            img_idx = 0
            for i, (tp, _, obj) in enumerate(relevant):
                # helper: run 提取
                def extract_runs(paras):
                    nonlocal img_idx
                    local_imgs = []
                    for para in paras:
                        for run in para.runs:
                            drawing = run._element.xpath('.//a:blip')
                            for blip in drawing:
                                rEmbed = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if rEmbed and rEmbed in doc.part.rels:
                                    rel = doc.part.rels[rEmbed]
                                    if "image" in rel.target_ref:
                                        img_idx += 1
                                        image_data = rel.target_part.blob
                                        img_filename = f"{doc_name}_advanced_img_{img_idx}.png"
                                        img_path = self.output_dir / img_filename
                                        with open(img_path, 'wb') as f:
                                            f.write(image_data)
                                        item = {
                                            'source_file': doc_path,
                                            'image_path': str(img_path),
                                            'image_type': 'embedded',
                                            'index': img_idx,
                                            'above_gqgqxx': True
                                        }
                                        local_imgs.append(item)
                    return local_imgs
                if tp == "para":
                    r = extract_runs([obj])
                    images.extend(r)
                    if found_bzdw_idx is not None and i < found_bzdw_idx:
                        bzdw_images.extend(r)
                else:  # cell
                    r = extract_runs(obj.paragraphs)
                    images.extend(r)
                    if found_bzdw_idx is not None and i < found_bzdw_idx:
                        bzdw_images.extend(r)
            return {
                'all_images_above_gqgqxx': images,
                'bzdw_images_count': len(bzdw_images),
                'bzdw_images': bzdw_images
            }
        except Exception as e:
            print(f"[extract_above_gqgqxx_and_check_bzdw_images] 出错: {e}")
            return {'all_images_above_gqgqxx': [], 'bzdw_images_count': 0, 'bzdw_images': []}


# ===================== PDF：关键词定位 + 区域裁剪 =====================
class PDFSignatureCropper:
    """
    在 PDF 中查找关键词（编制/校对/审核），将所在“行”扩展成一个矩形区域并裁剪渲染为图片。
    命中第一个就返回。
    """

    def __init__(self,
                 output_dir: str = "cropped_regions",
                 keywords: Optional[List[str]] = None,
                 padding: float = 20.0,
                 zoom: float = 2.0):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.keywords = keywords or ["编制", "校对", "审核"]
        self.padding = float(padding)
        self.zoom = float(zoom)

    @staticmethod
    def _union_rects(rects: List[fitz.Rect]) -> fitz.Rect:
        x0 = min(r.x0 for r in rects)
        y0 = min(r.y0 for r in rects)
        x1 = max(r.x1 for r in rects)
        y1 = max(r.y1 for r in rects)
        return fitz.Rect(x0, y0, x1, y1)

    @staticmethod
    def _overlap_1d(a0: float, a1: float, b0: float, b1: float) -> float:
        return max(0.0, min(a1, b1) - max(a0, b0))

    def _expand_to_row_band(self, page: fitz.Page, hit: fitz.Rect) -> fitz.Rect:
        """
        将命中的关键词矩形，扩展为与其“同一行”的文本块联合区域。
        """
        blocks = page.get_text("blocks")  # [(x0,y0,x1,y1, text, ...), ...]
        row_rects = []
        for b in blocks:
            x0, y0, x1, y1 = b[:4]
            br = fitz.Rect(x0, y0, x1, y1)
            # 与关键词在垂直方向有重叠，就认为在同一行/邻近行
            if self._overlap_1d(hit.y0, hit.y1, br.y0, br.y1) > 0:
                row_rects.append(br)

        region = self._union_rects(row_rects + [hit]) if row_rects else hit
        # 加 padding 并裁剪到页面
        region = fitz.Rect(
            max(page.rect.x0, region.x0 - self.padding),
            max(page.rect.y0, region.y0 - self.padding),
            min(page.rect.x1, region.x1 + self.padding),
            min(page.rect.y1, region.y1 + self.padding),
        )
        return region

    def find_and_crop_first(self, pdf_path: str) -> Optional[Dict]:
        """
        在 PDF 中寻找首个关键词出现区域，裁剪并渲染为图片。
        返回: dict{ image_path, page, region:[x0,y0,x1,y1] } 或 None
        """
        doc = fitz.open(pdf_path)
        try:
            for pno in range(len(doc)):
                page = doc[pno]
                hits = []
                for kw in self.keywords:
                    # quads=False -> 返回 Rect 列表，兼容性好
                    for r in page.search_for(kw, quads=False):
                        hits.append(fitz.Rect(r))
                if not hits:
                    continue
                # 读序：y0优先，其次x0
                hits.sort(key=lambda r: (round(r.y0, 2), round(r.x0, 2)))
                first_hit = hits[0]
                region = self._expand_to_row_band(page, first_hit)

                mat = fitz.Matrix(self.zoom, self.zoom)
                pix = page.get_pixmap(matrix=mat, clip=region, alpha=False)

                out_name = f"{Path(pdf_path).stem}_p{pno+1}_signature_region.png"
                out_path = self.output_dir / out_name
                pix.save(str(out_path))

                return {
                    "image_path": str(out_path),
                    "page": pno + 1,
                    "region": [float(region.x0), float(region.y0), float(region.x1), float(region.y1)]
                }
            return None
        finally:
            doc.close()


# ===================== VLM：手写痕迹检测 =====================
class HandwritingDetector:
    """手写痕迹检测器（保持你的调用方式）"""

    def __init__(self, api_config: Dict = None):
        self.api_config = api_config or self._get_default_config()

    def _get_default_config(self) -> Dict:
        return {
            "provider": "openai",
            "api_key": os.getenv("OPENAI_API_KEY", ""),
            "base_url": "https://api.openai.com/v1",
            "model": "gpt-4-vision-preview"
        }

    def encode_image(self, image_path: str) -> str:
        with open(image_path, "rb") as f:
            return base64.b64encode(f.read()).decode("utf-8")

    def detect_handwriting_in_signatures(self, image_path: str) -> Dict:
        if not self.api_config.get("api_key"):
            return {
                "error": "未配置VLM API密钥",
                "has_handwriting": False,
                "confidence": 0,
                "details": {},
            }

        try:
            base64_image = self.encode_image(image_path)
            prompt = """
请仔细分析这张图片，重点关注以下几个方面：

1. 查找图片中是否包含"编制"、"校对"、"审核"等字样的表格或签名栏
2. 检查这些区域是否有手写的名字、签名或其他手写痕迹
3. 区分打印文字和手写文字
回答只说有没有手写痕迹

"""
            if self.api_config["provider"] == "openai":
                return self._call_openai_vision(base64_image, prompt)
            else:
                return {"error": f"不支持的API提供商: {self.api_config['provider']}"}
        except Exception as e:
            return {"error": f"VLM分析出错: {str(e)}"}

    def _call_openai_vision(self, base64_image: str, prompt: str) -> Dict:
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.api_config['api_key']}"
        }
        payload = {
            "model": self.api_config["model"],
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                    ]
                }
            ],
            "max_tokens": 1000
        }
        resp = requests.post(
            f"{self.api_config['base_url']}/chat/completions",
            headers=headers,
            json=payload,
            timeout=60
        )
        if resp.status_code == 200:
            data = resp.json()
            content = data["choices"][0]["message"]["content"]
            try:
                return json.loads(content)
            except json.JSONDecodeError:
                return {
                    "error": "API返回非JSON格式",
                    "raw_response": content,
                    "has_handwriting": ("手写" in content) or ("签名" in content)
                }
        else:
            return {"error": f"API调用失败: {resp.status_code}", "message": resp.text}


# ===================== 处理器：合并 PDF 裁剪优先 + 首个命中即停 =====================
class DocumentProcessor:
    """文档处理主类"""

    def __init__(self,
                 output_dir: str = "results",
                 api_config: Dict = None,
                 stop_after_first: bool = True,
                 pdf_keywords: Optional[List[str]] = None,
                 pdf_padding: float = 20.0,
                 pdf_zoom: float = 2.0):
        self.extractor = DocumentImageExtractor(f"{output_dir}/images")
        self.detector = HandwritingDetector(api_config)
        self.cropper = PDFSignatureCropper(
            output_dir=f"{output_dir}/pdf_crops",
            keywords=pdf_keywords,
            padding=pdf_padding,
            zoom=pdf_zoom
        )
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.stop_after_first = stop_after_first

    def _is_signature_hit(self, result: Dict) -> bool:
        try:
            if result.get("has_signature_area") is True:
                return True
            if result.get("overall_assessment", {}).get("has_any_handwriting") is True:
                return True
            for area in result.get("signature_areas") or []:
                if area.get("has_handwriting") is True:
                    return True
            if result.get("has_handwriting") is True:
                return True
        except Exception:
            pass
        return False

    def _summarize_results(self, doc_path: Path, analysis_results: List[Dict], first_hit_path: Optional[str]) -> Dict:
        has_any_handwriting = any(r.get("_is_hit") for r in analysis_results)
        signature_findings = []
        errors = []
        for r in analysis_results:
            if "error" in r:
                errors.append(r["error"])
            if r.get("signature_areas"):
                signature_findings.extend(r["signature_areas"])
        return {
            "document": str(doc_path),
            "timestamp": str(pd.Timestamp.now()) if pd is not None else str(datetime.now()),
            "total_images_checked": len(analysis_results),
            "has_handwriting_signatures": has_any_handwriting,
            "signature_areas_found": len(signature_findings),
            "signature_details": signature_findings,
            "errors": errors,
            "detailed_results": analysis_results,
            "first_signature_image": first_hit_path
        }

    def process_pdf(self, pdf_path: str) -> Dict:
        """PDF 优先：关键词定位裁剪（首个），否则回退逐图"""
        doc_path = Path(pdf_path)
        print(f"正在处理 PDF：{doc_path.name}")

        analysis_results = []
        first_hit_path = None

        # 1) 关键词定位 + 裁剪（只拿到首个区域）
        crop = self.cropper.find_and_crop_first(pdf_path)
        if crop:
            print(f"已定位首个关键词区域：第 {crop['page']} 页 {crop['region']}")
            res = self.detector.detect_handwriting_in_signatures(crop["image_path"])
            res['image_info'] = {"source_file": pdf_path, "image_path": crop["image_path"],
                                 "image_type": "pdf_signature_region", "page": crop["page"]}
            res['_is_hit'] = self._is_signature_hit(res)
            analysis_results.append(res)
            first_hit_path = crop["image_path"]

            # 不论是否判定为有手写，都已完成“首个区域”的识别，直接汇总返回
            summary = self._summarize_results(doc_path, analysis_results, first_hit_path)
            self._save_summary(doc_path, summary)
            return summary

        # 2) 回退方案：提取页面图片并逐图检测（命中即停）
        print("未定位到关键词，回退到逐图检测...")
        images = self.extractor.extract_from_pdf(pdf_path)
        for img in images:
            res = self.detector.detect_handwriting_in_signatures(img["image_path"])
            res['image_info'] = img
            res['_is_hit'] = self._is_signature_hit(res)
            analysis_results.append(res)
            if res['_is_hit'] and self.stop_after_first:
                first_hit_path = img["image_path"]
                print(f"首个命中已找到，停止后续图片分析：{first_hit_path}")
                break

        summary = self._summarize_results(doc_path, analysis_results, first_hit_path)
        self._save_summary(doc_path, summary)
        return summary

    def extract_above_gqgqxx_and_check_bzdw_images_pdf(self, pdf_path: str, gqgqxx_regex: str, bzdw_regex: str):
        """
        PDF版本：
        1. 定位“高后果区基本信息”出现页及y坐标，只统计其上方所有图片；
        2. 在这个范围，再查找“编制单位”出现-y，统计其上方图片数。
        3. 返回{'all_images_above_gqgqxx': [...], 'bzdw_images_count': int, 'bzdw_images': [...]}。
        """
        import re
        pdf_doc = fitz.open(pdf_path)
        doc_name = Path(pdf_path).stem
        gqgqxx_prog = re.compile(gqgqxx_regex)
        bzdw_prog = re.compile(bzdw_regex)
        images = []
        bzdw_images = []
        found_gqgq_page = None
        found_gqgq_y = None
        found_bzdw_pg = None
        found_bzdw_y = None
        # 1. 遍历找高后果区基本信息（分界点）
        for pno in range(len(pdf_doc)):
            page = pdf_doc[pno]
            blocks = page.get_text('blocks')  # [(x0,y0,x1,y1, text,...)]
            for b in blocks:
                x0, y0, x1, y1, text, *_ = b
                if gqgqxx_prog.search(text):
                    found_gqgq_page, found_gqgq_y = pno, y0
                    break
            if found_gqgq_page is not None:
                break
        # 2. 在此分界点上方的所有图片全部提取
        img_idx = 0
        for pno in range(len(pdf_doc)):
            if found_gqgq_page is not None and pno > found_gqgq_page:
                break
            page = pdf_doc[pno]
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                pix = fitz.Pixmap(pdf_doc, xref)
                # 关联图片出现的块约束在分界点上方
                # 获取图片插入的大致y位置
                img_blocks = page.get_text('blocks')
                ranges = [(b[1], b[3]) for b in img_blocks]
                # 默认全部图片；若有分界点则按y裁
                max_y = found_gqgq_y if (found_gqgq_page == pno and found_gqgq_y is not None) else None
                # 不用精细y坐标，按页面判断即可
                if found_gqgq_page is not None:
                    if pno > found_gqgq_page or (pno == found_gqgq_page and max_y is not None):
                        continue
                img_idx += 1
                img_filename = f"{doc_name}_advanced_img_{img_idx}.png"
                img_path = self.output_dir / img_filename
                pix.save(str(img_path))
                images.append({'source_file': pdf_path, 'image_path': str(img_path), 'image_type': 'page_image', 'page': pno + 1, 'index': img_idx})
                pix = None
        # 3. 在这些页面/文本块内查找“编制单位”分界
        # 索引所有和图片属于同范围的文本（只考虑分界点上方页面）
        texts = []
        for pno in range(len(pdf_doc)):
            if found_gqgq_page is not None and pno > found_gqgq_page:
                break
            page = pdf_doc[pno]
            blocks = page.get_text('blocks')
            for b in blocks:
                x0, y0, x1, y1, text, *_ = b
                texts.append({'page': pno, 'y0': y0, 'y1': y1, 'text': text})
        # 找到“编制单位”出现的最大page和y0
        bzdw_ref = None
        for t in texts:
            if bzdw_prog.search(t['text']):
                bzdw_ref = t
                break
        # 统计“编制单位”上方的图片（按页面和y0小于分界即可）
        for img in images:
            # 仅做粗略分界（pdf图片精确到块级需ocr,这里页面优先）
            if bzdw_ref is not None:
                if img['page'] - 1 < bzdw_ref['page']:
                    bzdw_images.append(img)
                elif img['page'] - 1 == bzdw_ref['page']:
                    # 若同页，比y
                    continue
        pdf_doc.close()
        return {'all_images_above_gqgqxx': images, 'bzdw_images_count': len(bzdw_images), 'bzdw_images': bzdw_images}

    def process_pdf_cover_first(self, pdf_path: str) -> Dict:
        doc_path = Path(pdf_path)
        gqgqxx_regex = r'一[、,.．\\.]?高后果区基本信息'
        bzdw_regex = r'编制.?单位'
        result = self.extract_above_gqgqxx_and_check_bzdw_images_pdf(str(doc_path), gqgqxx_regex, bzdw_regex)
        all_imgs = result['all_images_above_gqgqxx']
        bzdw_count = result['bzdw_images_count']
        bzdw_imgs = result['bzdw_images']
        summary_reason = None
        first_hit_path = None
        analysis_results = []
        if bzdw_count == 3:
            summary = {
                "document": str(doc_path),
                "timestamp": str(pd.Timestamp.now()) if pd is not None else str(datetime.now()),
                "total_images_checked": 3,
                "has_handwriting_signatures": True,
                "signature_areas_found": 3,
                "signature_details": bzdw_imgs,
                "errors": [],
                "detailed_results": bzdw_imgs,
                "first_signature_image": bzdw_imgs[0]["image_path"] if bzdw_imgs else None,
                "reason": "‘一、高后果区基本信息’上方，‘编制单位’上方图片恰好3张，直接判定为封面手签，无需模型检测"
            }
            self._save_summary(doc_path, summary)
            return summary
        else:
            summary_reason = f"‘编制单位’上方图片数量为{bzdw_count}，未达到3张规则，对‘一、高后果区基本信息’上方全部图片进行手写检测"
            for img in all_imgs:
                res = self.detector.detect_handwriting_in_signatures(img["image_path"])
                res['image_info'] = img
                res['_is_hit'] = self._is_signature_hit(res)
                analysis_results.append(res)
                if res['_is_hit'] and self.stop_after_first:
                    first_hit_path = img["image_path"]
                    break
            summary = self._summarize_results(doc_path, analysis_results, first_hit_path)
            summary["reason"] = summary_reason
            self._save_summary(doc_path, summary)
            return summary

    def process_docx(self, docx_path: str) -> Dict:
        # 兼容旧接口，走封面图片优先策略
        return self.process_docx_cover_first(docx_path)

    def process_document(self, path: str) -> Dict:
        path = Path(path)
        if not path.exists():
            return {"error": f"文档不存在: {path}"}
        if path.suffix.lower() == ".pdf":
            return self.process_pdf_cover_first(str(path))
        elif path.suffix.lower() in [".docx", ".doc"]:
            return self.process_docx_cover_first(str(path))
        else:
            return {"error": f"不支持的文件格式: {path.suffix}"}

    def process_directory(self, dir_path: str, recursive: bool = True) -> Dict:
        dir_path = Path(dir_path)
        if not dir_path.exists():
            return {"error": f"目录不存在: {dir_path}"}
        patterns = ['*.pdf', '*.docx', '*.doc']
        files = []
        for pat in patterns:
            files.extend(dir_path.rglob(pat) if recursive else dir_path.glob(pat))

        print(f"找到 {len(files)} 个文档")
        results = []
        for f in files:
            try:
                results.append(self.process_document(str(f)))
            except Exception as e:
                results.append({"document": str(f), "error": f"处理失败: {e}"})

        batch = {
            "batch_processing": True,
            "directory": str(dir_path),
            "total_documents": len(files),
            "processed_successfully": len([r for r in results if "error" not in r]),
            "documents_with_handwriting": len([r for r in results if r.get("has_handwriting_signatures")]),
            "results": results
        }
        out = self.output_dir / f"batch_{dir_path.name}_summary.json"
        with open(out, "w", encoding="utf-8") as f:
            json.dump(batch, f, ensure_ascii=False, indent=2)
        return batch

    def _save_summary(self, doc_path: Path, summary: Dict):
        out = self.output_dir / f"{doc_path.stem}_handwriting_analysis.json"
        with open(out, "w", encoding="utf-8") as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)


# ===================== 对外导出：供 main.py 调用的消息构建函数 =====================
def build_vlm_messages_from_document(doc_path: str, first_only: bool = False) -> List[Dict]:
    """
    构建用于视觉大模型调用的消息数组：
    - PDF：优先按关键词裁剪出首个签字行区域；若无，则回退提取页面图片
    - DOCX/DOC：提取“‘一、高后果区基本信息’上方，且在‘编制单位’上方”的图片优先；若无则取上方所有
    返回格式与 `yingxiangtu.py` 保持一致：`[{"role":"user","content":[...]}]`
    """
    p = Path(doc_path)
    processor = DocumentProcessor(output_dir="handwriting_detection_results")
    contents: List[Dict] = []

    try:
        if p.suffix.lower() == ".pdf":
            crop = processor.cropper.find_and_crop_first(str(p))
            if crop:
                contents.append({"type": "image", "image": crop["image_path"]})
            if not contents:
                for it in processor.extractor.extract_from_pdf(str(p)):
                    contents.append({"type": "image", "image": it.get("image_path")})
        elif p.suffix.lower() in [".docx", ".doc"]:
            try:
                rr = processor.extractor.extract_above_gqgqxx_and_check_bzdw_images(
                    str(p),
                    gqgqxx_regex=r'一[、,.．\\.]?高后果区基本信息',
                    bzdw_regex=r'编制.?单位'
                )
                picks = rr.get("bzdw_images") or rr.get("all_images_above_gqgqxx") or []
                for it in picks:
                    contents.append({"type": "image", "image": it.get("image_path")})
            except Exception:
                pass
    except Exception:
        # 若发生异常，保持返回空图像，仅附带文本提示
        pass

    if first_only and contents:
        contents = [contents[0]]

    contents.append({
        "type": "text",
        "text": (
            "请判断图中‘编制/校对/审核’签字栏是否存在手写签名或手写痕迹。"
            "只需回答是否存在手写痕迹：是/否。"
        )
    })
    return [{"role": "user", "content": contents}]


# ===================== CLI =====================
def parse_args():
    ap = argparse.ArgumentParser(description="签字检测：输出消息数组或检测结果数组（参照 yingxiangtu.py 风格）")
    ap.add_argument("--input", required=False, default=None, help="输入文件或目录（PDF/DOCX/DOC）。不提供则默认扫描 word-master/")
    ap.add_argument("--mode", choices=["messages", "detect", "auto"], default="messages",
                    help="messages: 仅构建VLM消息数组；detect: 构建后调用模型返回检测结果数组；auto: 保持旧流程自动处理并输出汇总")
    ap.add_argument("--first_only", action="store_true", help="仅取首个命中区域/首张相关图片")
    ap.add_argument("--recursive", action="store_true", help="目录递归处理（仅 auto 模式生效）")
    ap.add_argument("--output_dir", default="handwriting_detection_results", help="输出目录")
    ap.add_argument("--keywords", default="编制,校对,审核", help="PDF关键词，逗号分隔")
    ap.add_argument("--padding", type=float, default=20.0, help="PDF裁剪padding(px)")
    ap.add_argument("--zoom", type=float, default=2.0, help="PDF渲染缩放倍数")
    ap.add_argument("--stop_after_first", action="store_true", help="命中首个即停（默认即停）")
    return ap.parse_args()


def main():
    args = parse_args()
    keywords = [s.strip() for s in args.keywords.split(",") if s.strip()] if args.keywords else []
    os.makedirs(args.output_dir, exist_ok=True)
    processor = DocumentProcessor(
        output_dir=args.output_dir,
        api_config={
            "provider": "openai",
            "api_key": os.getenv("OPENAI_API_KEY", ""),
            "base_url": "https://api.openai.com/v1",
            "model": "gpt-4-vision-preview"
        },
        stop_after_first=True if args.stop_after_first or True else False,
        pdf_keywords=keywords if keywords else None,
        pdf_padding=args.padding,
        pdf_zoom=args.zoom
    )

    def build_signature_messages(doc_path: Path, first_only: bool = False) -> List[Dict]:
        prompt_text = (
            "请判断图中‘编制/校对/审核’签字栏是否存在手写签名或手写痕迹。"
            "只需回答是否存在手写痕迹：是/否。"
        )
        contents: List[Dict] = []
        try:
            if doc_path.suffix.lower() == ".pdf":
                crop = processor.cropper.find_and_crop_first(str(doc_path))
                if crop:
                    contents.append({"type": "image", "image": crop["image_path"]})
            elif doc_path.suffix.lower() in [".docx", ".doc"]:
                try:
                    rr = processor.extractor.extract_above_gqgqxx_and_check_bzdw_images(
                        str(doc_path),
                        gqgqxx_regex=r'一[、,.．\\.]?高后果区基本信息',
                        bzdw_regex=r'编制.?单位'
                    )
                    picks = rr.get("bzdw_images") or rr.get("all_images_above_gqgqxx") or []
                    for it in picks:
                        contents.append({"type": "image", "image": it.get("image_path")})
                except Exception:
                    pass
        except Exception:
            pass

        if not contents and doc_path.suffix.lower() == ".pdf":
            for it in processor.extractor.extract_from_pdf(str(doc_path)):
                contents.append({"type": "image", "image": it.get("image_path")})

        if first_only and contents:
            contents = [contents[0]]

        contents.append({"type": "text", "text": prompt_text})
        return [{"role": "user", "content": contents}]

    def detect_from_messages(messages: List[Dict]) -> List[Dict]:
        results: List[Dict] = []
        try:
            for msg in messages:
                for c in (msg.get("content") or []):
                    if c.get("type") == "image" and c.get("image"):
                        r = processor.detector.detect_handwriting_in_signatures(c.get("image"))
                        r["image_path"] = c.get("image")
                        results.append(r)
        except Exception as e:
            results.append({"error": f"检测失败: {e}"})
        return results

    def _is_supported_file(p: Path) -> bool:
        return p.suffix.lower() in [".pdf", ".docx", ".doc"] and p.exists() and not p.name.startswith("~$")

    def _auto_discover_single() -> Optional[Path]:
        cwd = Path.cwd()
        wm = cwd / "word-master"
        # 优先 word-master 下
        if wm.exists():
            for pat in ("*.pdf", "*.docx", "*.doc"):
                for p in wm.glob(pat):
                    if _is_supported_file(p):
                        return p
        # 其次当前目录
        for pat in ("*.pdf", "*.docx", "*.doc"):
            for p in cwd.glob(pat):
                if _is_supported_file(p):
                    return p
        # 最后全局搜索首个
        for p in cwd.rglob("*.pdf"):
            if _is_supported_file(p):
                return p
        for p in cwd.rglob("*.docx"):
            if _is_supported_file(p):
                return p
        for p in cwd.rglob("*.doc"):
            if _is_supported_file(p):
                return p
        return None

    # 解析输入：仅处理一个文件
    chosen: Optional[Path] = None
    if args.input:
        in_path = Path(args.input)
        if not in_path.exists():
            print(json.dumps({"error": f"路径不存在：{in_path}"}, ensure_ascii=False))
            return
        if in_path.is_dir():
            # 目录时自动选首个支持的文件（非递归优先）
            for pat in ("*.pdf", "*.docx", "*.doc"):
                for p in in_path.glob(pat):
                    if _is_supported_file(p):
                        chosen = p
                        break
                if chosen:
                    break
            # 若非递归没找到，再递归
            if not chosen:
                for pat in ("*.pdf", "*.docx", "*.doc"):
                    for p in in_path.rglob(pat):
                        if _is_supported_file(p):
                            chosen = p
                            break
                    if chosen:
                        break
        else:
            chosen = in_path if _is_supported_file(in_path) else None
    else:
        # 优先使用样例文件，其次自动发现
        sample = Path(SAMPLE_FILE_PATH) if 'SAMPLE_FILE_PATH' in globals() else None
        if SAMPLE_FILE_PATH and sample and sample.exists() and sample.is_file():
            chosen = sample
        else:
            chosen = _auto_discover_single()

    if not chosen:
        print(json.dumps({"error": "未找到可处理的文件，请使用 --input 指定 PDF/DOCX 文件"}, ensure_ascii=False))
        return

    # 三种模式：对单文件操作
    if args.mode == "auto":
        res = processor.process_document(str(chosen))
        print(json.dumps(res, ensure_ascii=False, indent=2))
        return

    msgs = build_signature_messages(chosen, first_only=args.first_only)
    if args.mode == "messages":
        print(json.dumps(msgs, ensure_ascii=False, indent=2))
    else:
        print(json.dumps(detect_from_messages(msgs), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
