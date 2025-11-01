import json
import time
import re
import os
import fitz
import datetime
from docx import Document

# ========================== 1. 配置区 ==========================
API_KEY = "BearerCubJNNWPCQjooxiEbLGB:IgSlCSEPDRuFXHJlXlMw" 
URL = "https://spark-api-open.xf-yun.com/v2/chat/completions"
INPUT_DOCUMENT_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/豫洛阳-兰郑长干线-CPY-0790-BFGDGS-ZZSYQFGS.docx"
OUTPUT_JSON_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/standard_compliance.json"

# 仅处理第一个语义块以快速预览效果（参考 template_compliance_check.py）
PROCESS_FIRST_CHUNK_ONLY = False

STANDARD_HCA_RULES = (
    "6.1.1.1 管道经过区域符合表1识别项中任何一条的为高后果区。\n"\
    "\n输油管道高后果区识别分级：\n"\
    "- a）管道中心线两侧各200m范围内，任意划分成长度为2km并能包括最大聚居户数的若干地段，四层及四层以上楼房（不计地下室层数）普遍集中、交通频繁、地下设施多的区段，分级为Ⅲ级。\n"\
    "- b）管道中心线两侧200m范围内，任意划分2km长度并能包括最大聚居户数的若干地段，户数在100户或以上的区段，包括市郊居住区、商业区、工业区、发展区以及不够四级地区条件的人口稠密区，分级为Ⅱ级。\n"\
    "- c）管道两侧各200m内有聚居户数在50户或以上的村庄、乡镇等，分级为Ⅱ级。\n"\
    "- d）管道两侧各50m内有高速公路、国道、省道、铁路及易燃易爆场所等，分级为Ⅰ级。\n"\
    "- e）管道两侧各200m内有湿地、森林、河口等国家自然保护地区，分级为Ⅰ级。\n"\
    "- f）管道两侧各200m内有水源、河流、大中型水库，分级为Ⅲ级。\n"\
    "\n6.1.1.2 识别高后果区时，高后果区边界设定为距离最近一幢建筑物外边缘200m。\n"\
    "\n6.1.1.3 高后果区分为三级，Ⅰ级代表最小的严重程度，Ⅲ级代表最大的严重程度。\n"\
    "\n6.1.2 输气管道高后果区\n"\
    "\n6.1.2.1 管道经过区域符合表2识别项中任何一条的为高后果区。\n"\
    "\n输气管道高后果区识别分级：\n"\
    "- a）管道经过的四级地区，地区等级按照GB 50251中相关规定执行，分级为Ⅲ级。\n"\
    "- b）管道经过的三级地区，分级为Ⅱ级。\n"\
    "- c）如管径大于762mm，并且最大允许操作压力大于6.9MPa，其天然气管道潜在影响区域内有特定场所的区域，潜在影响半径按式(1)计算，分级为Ⅱ级。\n"\
    "- d）如管径小于273mm，并且最大允许操作压力小于1.6MPa，其天然气管道潜在影响区域内有特定场所的区域，潜在影响半径按式(1)计算，分级为Ⅰ级。\n"\
    "- e）其他管道两侧各200m内有特定场所的区域，分级为Ⅰ级。\n"\
    "- f）除三级、四级地区外，管道两侧各200m内有加油站、油库等易燃易爆场所，分级为Ⅱ级。"
)


# ========================== 2. 功能函数区 ==========================

def read_document(file_path):
    """读取Word (.docx, .doc) 或 PDF (.pdf) 文档，返回纯文本内容"""
    full_text = []
    file_path_lower = file_path.lower()
    try:
        if file_path_lower.endswith('.docx'):
            from docx import Document
            doc = Document(file_path)
            for para in doc.paragraphs:
                full_text.append(para.text)
            return "\n".join(full_text)
            
        elif file_path_lower.endswith('.doc'):
            try:
                import subprocess
                output_dir = os.path.dirname(file_path)
                subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'txt:Text', '--outdir', output_dir, file_path],
                    check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120
                )
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                txt_file_path = os.path.join(output_dir, f"{base_name}.txt")
                if os.path.exists(txt_file_path):
                    with open(txt_file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    os.remove(txt_file_path)
                    return content
                else:
                    return f"错误: LibreOffice 转换后未找到文本文件 {txt_file_path}"
            except FileNotFoundError:
                return "错误: 'libreoffice' 命令未找到。请先安装 (sudo apt-get install libreoffice-writer)。"
            except subprocess.TimeoutExpired:
                return f"错误: 使用 LibreOffice 转换文件超时 ({file_path})。"
            except subprocess.CalledProcessError as e:
                error_message = e.stderr.decode('utf-8', errors='ignore')
                return f"错误: 使用 LibreOffice 读取 .doc 文件失败。\n详细错误: {error_message}"
                
        elif file_path_lower.endswith('.pdf'):
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                full_text.append(page.get_text())
            doc.close()
            return "\n".join(full_text)
        else:
            return f"错误: 不支持的文件格式 {file_path}"
    except Exception as e:
        return f"错误: 读取文件失败 {file_path} - {str(e)}"

def semantic_chunker(text, max_chars=250, overlap=50):
    """将长文本进行自然的语义切块"""
    text = re.sub(r'\n\s*\n', '\n', text)
    paragraphs = text.split('\n')
    chunks = []
    current_chunk = ""
    
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        if len(current_chunk) + len(p) + 1 <= max_chars:
            current_chunk += p + "\n"
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = p + "\n"
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

 

# ========================== 类型判定（不生成负样本） ==========================
def _clean(s: str) -> str:
    return (s or "").strip()


def _extract_value_from_table_cell(doc: Document, keyword: str, exclude_keywords: list[str] = None, max_search_range: int = 3) -> str | None:
    """
    从表格中提取指定关键词对应的值。
    查找逻辑：
    1. 在当前单元格文本中查找关键词（支持"关键词: 值"格式）
    2. 在当前行的右侧单元格查找
    3. 在当前列的下方单元格查找
    
    Args:
        doc: Document对象
        keyword: 要查找的关键词（如"高后果区类型"、"高后果区长度"等）
        exclude_keywords: 需要排除的关键词列表（用于避免"类型"匹配到"高后果区类型"）
        max_search_range: 向下搜索的最大行数
    
    Returns:
        找到的值字符串（已去重），未找到返回None
    """
    keyword_normalized = keyword.replace(" ", "").replace("　", "")
    exclude_normalized = set()
    if exclude_keywords:
        exclude_normalized = {k.replace(" ", "").replace("　", "") for k in exclude_keywords}
    
    # 确保当前关键词不会被排除
    exclude_normalized.discard(keyword_normalized)
    
    for tbl in doc.tables:
        rows = list(getattr(tbl, "rows", []))
        for r_idx, row in enumerate(rows):
            cells = list(getattr(row, "cells", []))
            for c_idx, cell in enumerate(cells):
                raw = _clean(cell.text)
                raw_normalized = raw.replace(" ", "").replace("　", "")
                
                # 检查是否包含关键词，但排除包含其他更具体关键词的情况
                if keyword_normalized in raw_normalized:
                    # 如果包含需要排除的关键词（更具体的关键词），跳过
                    # 例如：当搜索"类型"时，如果单元格包含"高后果区类型"，应该跳过
                    should_exclude = False
                    for excl_kw in exclude_normalized:
                        if excl_kw in raw_normalized:
                            # 如果排除关键词比当前关键词更长或更具体，则排除
                            if len(excl_kw) > len(keyword_normalized) or excl_kw != keyword_normalized:
                                should_exclude = True
                                break
                    if should_exclude:
                        continue
                    
                    # 方式1: 当前单元格内包含"关键词: 值"格式
                    patterns = [
                        rf"{re.escape(keyword)}[:：]\s*(.+)",
                        rf"{re.escape(keyword)}\s+[:：]\s*(.+)",
                    ]
                    for pattern in patterns:
                        m = re.search(pattern, raw)
                        if m:
                            value = _clean(m.group(1))
                            if value and value != keyword:
                                return _deduplicate_value(value)
                    
                    # 方式2: 查找右侧单元格（仅取第一个有效值）
                    for c in cells[c_idx + 1:]:
                        text = _clean(c.text)
                        if text and text != keyword and text not in keyword_normalized:
                            # 检查是否包含排除关键词
                            text_normalized = text.replace(" ", "").replace("　", "")
                            if not any(excl_kw in text_normalized for excl_kw in exclude_normalized):
                                return _deduplicate_value(text)
                    
                    # 方式3: 查找下方单元格（仅取第一个有效值）
                    for k in range(1, max_search_range + 1):
                        rr = r_idx + k
                        if rr < len(rows):
                            c2 = rows[rr].cells
                            if c_idx < len(c2):
                                v = _clean(c2[c_idx].text)
                                if v and v != keyword and v not in keyword_normalized:
                                    # 检查是否包含排除关键词
                                    v_normalized = v.replace(" ", "").replace("　", "")
                                    if not any(excl_kw in v_normalized for excl_kw in exclude_normalized):
                                        return _deduplicate_value(v)
    return None


def _deduplicate_value(value: str) -> str:
    """
    对提取的值进行去重处理。
    例如："人员密集型、人员密集型" -> "人员密集型"
    """
    if not value:
        return value
    
    # 如果包含分隔符（、，,等），分割后去重
    separators = ["、", ",", "，", "|", "/"]
    parts = [value]
    for sep in separators:
        if sep in parts[0]:
            new_parts = []
            for part in parts:
                new_parts.extend([p.strip() for p in part.split(sep)])
            parts = new_parts
            break
    
    # 去重并保留顺序
    seen = set()
    unique_parts = []
    for part in parts:
        part_clean = part.strip()
        if part_clean and part_clean not in seen:
            seen.add(part_clean)
            unique_parts.append(part_clean)
    
    # 重新组合
    return "、".join(unique_parts) if len(unique_parts) > 1 else (unique_parts[0] if unique_parts else value)


def _extract_hca_info_from_tables(doc: Document) -> dict:
    """
    从"高后果区基本信息表"中提取所有相关信息。
    返回字典，包含：高后果区类型、高后果区长度(m)、高后果区等级、识别项
    """
    result = {
        "高后果区类型": None,
        "高后果区长度": None,
        "高后果区等级": None,
        "识别项": None,
    }
    
    # 定义要提取的字段关键词（支持多种可能的表述）
    # 注意：按优先级排序，更具体的放在前面，用于exclude_keywords避免重复匹配
    field_keywords = {
        "高后果区类型": {
            "keywords": ["高后果区类型", "类型"],
            "exclude": ["高后果区长度", "高后果区等级", "识别项"]  # 避免误匹配
        },
        "高后果区长度": {
            "keywords": ["高后果区长度(m)", "高后果区长度", "长度(m)", "长度（m）", "长度"],
            "exclude": ["高后果区类型", "高后果区等级", "识别项", "类型"]  # "长度"不能匹配到"高后果区长度"
        },
        "高后果区等级": {
            "keywords": ["高后果区等级", "等级", "分级"],
            "exclude": ["高后果区类型", "高后果区长度", "识别项", "类型", "长度"]
        },
        "识别项": {
            "keywords": ["识别项", "识别依据"],
            "exclude": ["高后果区类型", "高后果区长度", "高后果区等级", "类型", "长度", "等级"]
        },
    }
    
    # 已提取的值集合，用于避免不同字段提取到相同的值
    extracted_values = set()
    
    # 提取每个字段（按优先级顺序）
    for field_name, config in field_keywords.items():
        keywords = config["keywords"]
        exclude = config.get("exclude", [])
        
        for keyword in keywords:
            value = _extract_value_from_table_cell(doc, keyword, exclude_keywords=exclude)
            if value:
                # 检查是否与其他已提取的值重复（规范化后比较）
                value_normalized = value.replace(" ", "").replace("　", "")
                if value_normalized not in extracted_values:
                    result[field_name] = value
                    extracted_values.add(value_normalized)
                    break
                elif not result[field_name]:  # 如果当前字段还没有值，即使重复也使用（可能是唯一值）
                    result[field_name] = value
                    break
    
    return result


def build_standard_compliance_from_docx(doc_path: str) -> str:
    """
    从"高后果区基本信息表"中提取相关信息。
    返回格式化的字符串，包含：高后果区类型、高后果区长度(m)、高后果区等级、识别项
    """
    try:
        if not os.path.exists(doc_path):
            return f"❌ 找不到文档：{doc_path}"
        doc = Document(doc_path)
        info = _extract_hca_info_from_tables(doc)
        
        # 格式化输出
        parts = []
        for key, value in info.items():
            display_key = key if key == "高后果区类型" else key.replace("高后果区", "")
            if key == "高后果区长度" and value:
                display_key = "高后果区长度(m)"
            parts.append(f"{display_key}：{value or '未识别'}")
        
        return "\n".join(parts)
    except Exception as e:
        return f"❌ 处理失败: {e}"

def build_vlm_messages_from_docx(doc_path: str, first_only: bool = False):
    """
    从 DOCX 文件中提取高后果区信息，生成用于验证标准符合性的 VLM 消息。
    
    返回格式与 `inference/model.py` 中 `make_messages` 一致（但不包含图片）：
    [
      {
        "role": "user",
        "content": [
          {"type": "text", "text": "...验证提示词..."}
        ]
      }
    ]
    """
    try:
        if not os.path.exists(doc_path):
            return [{
                "role": "user",
                "content": [{"type": "text", "text": f"❌ 找不到文档：{doc_path}"}]
            }]
        
        doc = Document(doc_path)
        extracted_info = _extract_hca_info_from_tables(doc)
        
        # 构建提取信息摘要
        info_summary = []
        for key, value in extracted_info.items():
            display_key = key if key == "高后果区类型" else key.replace("高后果区", "")
            if key == "高后果区长度" and value:
                display_key = "高后果区长度(m)"
            info_summary.append(f"{display_key}：{value or '未识别'}")
        extracted_info_text = "\n".join(info_summary)
        
        # 构建验证 prompt
        prompt_text = (
            "请基于以下高后果区识别标准和提取的信息进行核查并严格按以下要求作答：\n\n"
            "【标准规则】\n"
            f"{STANDARD_HCA_RULES}\n\n"
            "【从文档中提取的高后果区信息】\n"
            f"{extracted_info_text}\n\n"
            "【验证要求】\n"
            "请逐一核查以下内容：\n"
            "1) 提取的高后果区等级（Ⅰ级/Ⅱ级/Ⅲ级）是否符合标准规则中的分级要求\n"
            "2) 提取的识别项是否与标准规则中的识别项对应\n"
            "【输出格式】\n"
            "请先给出'是否符合标准：是/否'的总体结论；\n"
            "如不符合，请逐项指出问题：\n"
            "- 具体哪些信息不符合标准规则\n"
            "- 正确的值应该是什么\n"
            "- 给出具体的改正建议\n"
        )
        
        return [{
            "role": "user",
            "content": [
                {"type": "text", "text": prompt_text}
            ]
        }]
        
    except Exception as e:
        return [{
            "role": "user",
            "content": [{"type": "text", "text": f"❌ 处理失败: {e}"}]
        }]


# ========================== 主程序 ==========================
if __name__ == '__main__':
    # 测试提取功能
    result = build_standard_compliance_from_docx(INPUT_DOCUMENT_PATH)
    print("=" * 60)
    print("提取的高后果区信息：")
    print(result)
    print("=" * 60)
    
    # 测试 VLM 消息生成
    vlm_messages = build_vlm_messages_from_docx(INPUT_DOCUMENT_PATH)
    print("\n生成的 VLM 消息格式：")
    print(json.dumps(vlm_messages, ensure_ascii=False, indent=2))
    print("=" * 60)
