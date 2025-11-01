import re
import os
import argparse
from pathlib import Path
from docx import Document

# ========================== 1. 配置区 ==========================
INPUT_DOCUMENT_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/豫郑州-洛驻线-CPY-0190-BFGDGS-ZZSYQFGS.docx"

# ========================== 2. 解析 Word ==========================
def parse_docx_sections(doc_path):
    """解析 Word 文档，识别主要表格区段"""
    doc = Document(doc_path)
    sections = {"basic_info": "", "risk_table": "", "prevention": "", "full_text": ""}
    current_section = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        sections["full_text"] += text + "\n"

        if re.search(r'高后果区基本信息|基本信息表', text):
            current_section = "basic_info"
        elif re.search(r'高后果区风险评价结果表|风险评价结果表|风险评价', text):
            current_section = "risk_table"
        elif re.search(r'人防措施|防护措施', text):
            current_section = "prevention"

        if current_section:
            sections[current_section] += text + "\n"

    # 合并表格文本
    for table in doc.tables:
        table_text = "\n".join(["  ".join([cell.text.strip() for cell in row.cells]) for row in table.rows])
        if re.search(r'高后果区基本信息|基本信息表', table_text):
            sections["basic_info"] += "\n" + table_text
        elif re.search(r'高后果区风险评价结果表|风险评价结果表|风险评价|已采取的控制措施|控制措施', table_text):
            sections["risk_table"] += "\n" + table_text
        elif re.search(r'人防|防护', table_text):
            sections["prevention"] += "\n" + table_text

    return sections

# ========================== 3. 信息提取 ==========================
def extract_section_leaders(sections):
    """提取区段长信息（支持提取多个，包含职位名称）"""
    leader_pattern = re.compile(
        r'(专职区（段）长姓名|专职区\(段\)长姓名|专职区长姓名|专职段长姓名|专职区段长姓名|区长姓名|专职区（段）长|专职区\(段\)长|专职区长|专职段长|专职区段长)[：:\s]*([\u4e00-\u9fa5·]{2,8})'
    )
    prevention_matches = re.findall(leader_pattern, sections.get("prevention", ""))
    basic_info_matches = re.findall(leader_pattern, sections.get("basic_info", ""))
    risk_table_matches = re.findall(leader_pattern, sections.get("risk_table", ""))
    
    # 过滤掉不完整的数据
    prevention_leaders = []
    for title, name in prevention_matches:
        cleaned = (name or '').strip()
        if not cleaned:
            continue
        if cleaned in {"姓名", "名字"}:
            continue
        if any(char in cleaned for char in ['等', '度', 'm', ')', '）']):
            continue
        prevention_leaders.append(f"{title}：{cleaned}")
    
    basic_info_leaders = []
    for title, name in basic_info_matches:
        cleaned = (name or '').strip()
        if not cleaned:
            continue
        if cleaned in {"姓名", "名字"}:
            continue
        if any(char in cleaned for char in ['等', '度', 'm', ')', '）']):
            continue
        basic_info_leaders.append(f"{title}：{cleaned}")
    
    risk_table_leaders = []
    for title, name in risk_table_matches:
        cleaned = (name or '').strip()
        if not cleaned:
            continue
        if cleaned in {"姓名", "名字"}:
            continue
        if any(char in cleaned for char in ['等', '度', 'm', ')', '）']):
            continue
        risk_table_leaders.append(f"{title}：{cleaned}")
    
    # 去重并保持顺序
    prevention_leaders = list(dict.fromkeys(prevention_leaders))
    basic_info_leaders = list(dict.fromkeys(basic_info_leaders))
    risk_table_leaders = list(dict.fromkeys(risk_table_leaders))
    
    return {
        "from_prevention_measures": prevention_leaders,
        "from_basic_info_table": basic_info_leaders,
        "from_risk_table": risk_table_leaders
    }

def extract_locations(sections):
    """提取位置信息（支持提取多个，包含位置类型）"""
    loc_pattern = re.compile(
        r'(位置|位于|行政区划|地址|地理位置|高后果区起点|高后果区终点)[：:\s]*([^\s，,。\n]{3,30})'
    )
    basic_info_matches = re.findall(loc_pattern, sections.get("basic_info", ""))
    risk_table_matches = re.findall(loc_pattern, sections.get("risk_table", ""))
    
    # 组合位置类型和具体位置
    basic_info_locations = [f"{loc_type}：{location}" for loc_type, location in basic_info_matches]
    risk_table_locations = [f"{loc_type}：{location}" for loc_type, location in risk_table_matches]
    
    # 去重并保持顺序
    basic_info_locations = list(dict.fromkeys(basic_info_locations))
    risk_table_locations = list(dict.fromkeys(risk_table_locations))
    
    return {
        "from_basic_info": basic_info_locations,
        "from_risk_table": risk_table_locations
    }

def extract_hca_numbers_legacy(sections):
    """提取高后果区编号（旧方法，从sections中提取）"""
    # 改进后的正则模式（兼容省份/线路/支线等多样写法）
    number_patterns = [
        # 通用完整格式，例如：
        r'[豫Y][\u4e00-\u9fa5A-Z0-9]+[-—_－]?[A-Z\u4e00-\u9fa5]*[-—_－]?(CPY[-_—]?\d{3,4}(?:[-_A-Z0-9]+)?)',

        # 简写或残缺格式（仅编号部分）
        r'(CPY[-_—]?\d{3,4}(?:[-_A-Z0-9]+)?)'
    ]

    found = set()
    nums = []
    text = sections.get("full_text", "")

    for p in number_patterns:
        for match in re.finditer(p, text):
            num = match.group(1) if match.lastindex else match.group(0)
            if num not in found:
                found.add(num)
                nums.append(num)

    return nums

def extract_hca_numbers(doc_path):
    """
    从指定Word文档中提取封面编号和高后果区基本信息表中的编号
    :param doc_path: Word文档完整路径
    :return: 提取结果字典，包含"封面编号"和"高后果区基本信息表编号"
    """
    result = {
        "封面编号": "未找到",
        "高后果区基本信息表编号": "未找到"
    }
    
    # 验证文件路径
    if not os.path.exists(doc_path):
        return result
    
    # 读取文档
    try:
        doc = Document(doc_path)
    except Exception as e:
        import traceback
        print(f"读取文档失败：{str(e)}")
        if False:  # 可以改为 True 启用调试
            traceback.print_exc()
        return result
    
    # 1. 提取封面编号（正则表达式匹配）
    pattern_cover = re.compile(r'(编号：|封面编号：)\s*([\u4e00-\u9fa5A-Za-z0-9-]+)')
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if para_text:  # 跳过空段落
            match = pattern_cover.search(para_text)
            if match:
                result["封面编号"] = match.group(2)
                break  # 找到后停止遍历段落
    
    # 2. 提取高后果区基本信息表中的编号
    pattern_table = re.compile(r'高后果区编号\s*[:：]?\s*([\u4e00-\u9fa5A-Za-z0-9-]+)')
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join([cell.text.strip() for cell in row.cells])
            if row_text:
                match = pattern_table.search(row_text)
                if match:
                    result["高后果区基本信息表编号"] = match.group(1)
                    return result  # 找到后直接返回
    
    return result

# 读取 w:t，包含文本框里的文字（不要用 xpath 的 namespaces，兼容性更好）
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _all_wt_text(elem):
    """读取元素内所有 w:t 节点的文本，包括文本框中的文字"""
    return "".join((t.text or "") for t in elem.iter(f"{{{W_NS}}}t"))

def _cell_text(cell):
    """获取单元格的完整文本（包括文本框），带上换行便于正则分段"""
    return _all_wt_text(cell._tc).replace("\r", "\n").strip()

def _norm(s: str) -> str:
    """标准化括号"""
    return s.replace("(", "（").replace(")", "）")

def extract_control_measures(doc_path, debug=False):
    """从风险评价表中提取【已采取的控制措施】里的【1 人防措施】文本"""
    try:
        doc = Document(doc_path)
        results = []
        debug_info = []

        for tbl_idx, tbl in enumerate(doc.tables):
            # 先把整张表拼成文本（用 _cell_text，能读到文本框），做一个快速过滤
            tbl_text = "\n".join("  ".join(_cell_text(c) for c in row.cells) for row in tbl.rows)
            is_risk_table = re.search(r'(风险评价结果|风险评价结果表|风险评价|计划/控制措施|已采取的控制措施|控制措施)', tbl_text)
            
            if debug:
                debug_info.append(f"表格 {tbl_idx}: 是否风险评价表={bool(is_risk_table)}, 表格文本长度={len(tbl_text)}")
            
            if not is_risk_table:
                continue

            # 再逐单元格查找"已采取的控制措施 … 1 人防措施 …"
            for row_idx, row in enumerate(tbl.rows):
                for col_idx, cell in enumerate(row.cells):
                    s_raw = _cell_text(cell)
                    s = _norm(s_raw)
                    
                    # 检查是否包含控制措施相关内容
                    has_control_keyword = re.search(r'(已采取的控制措施|计划/控制措施|控制措施|人防措施)', s)
                    if not has_control_keyword:
                        continue
                    
                    if debug:
                        debug_info.append(f"  单元格[{row_idx},{col_idx}]: 长度={len(s)}, 前100字符={repr(s[:100])}")

                    # 策略1：主匹配：在"已采取的控制措施"块里切出"1 人防措施"
                    m = re.search(
                        r'已采取的控制措施[\s\S]{0,800}?'
                        r'(?:1[）\.\、]?\s*)?人防措施[：:\s]*([\s\S]{0,2000}?)'
                        r'(?=\n\s*(?:2[）\.\、]?\s*|拟采取|计划/控制措施|控制措施|$|\n\n))',
                        s, flags=re.S
                    )
                    if m:
                        extracted = m.group(1).strip()
                        if extracted and len(extracted) > 5:
                            results.append(extracted)
                            if debug:
                                debug_info.append(f"    ✓ 策略1提取成功: {len(extracted)}字符")
                            continue

                    # 策略2：兜底：单元格里只有"1 人防措施 …"这一块
                    m2 = re.search(
                        r'(?:1[）\.\、]?\s*)?人防措施[：:\s]*([\s\S]{0,2000}?)'
                        r'(?=\n\s*(?:2[）\.\、]?\s*|拟采取|计划/控制措施|控制措施|$|\n\n))',
                        s, flags=re.S
                    )
                    if m2:
                        extracted = m2.group(1).strip()
                        if extracted and len(extracted) > 5:
                            results.append(extracted)
                            if debug:
                                debug_info.append(f"    ✓ 策略2提取成功: {len(extracted)}字符")
                            continue
                    
                    # 策略3：如果单元格包含"已采取的控制措施"，直接提取整个单元格内容（截取相关部分）
                    if re.search(r'已采取的控制措施', s):
                        # 尝试提取"已采取的控制措施"之后的所有内容
                        m3 = re.search(r'已采取的控制措施[：:\s]*([\s\S]{0,3000}?)', s, flags=re.S)
                        if m3:
                            extracted = m3.group(1).strip()
                            # 过滤掉明显不是内容的部分
                            if extracted and len(extracted) > 10 and not re.match(r'^\d+$', extracted):
                                results.append(extracted)
                                if debug:
                                    debug_info.append(f"    ✓ 策略3提取成功: {len(extracted)}字符")
                    
                    # 策略4：如果单元格包含"人防措施"，提取人防措施相关内容
                    elif re.search(r'人防措施', s) and len(s) > 20:
                        # 提取包含"人防措施"的段落
                        m4 = re.search(r'([^\n]*人防措施[：:\s]*[^\n]{10,2000}?)', s, flags=re.S)
                        if m4:
                            extracted = m4.group(1).strip()
                            if extracted and len(extracted) > 10:
                                results.append(extracted)
                                if debug:
                                    debug_info.append(f"    ✓ 策略4提取成功: {len(extracted)}字符")
                    
                    # 策略5：如果单元格很长且包含相关关键词，直接返回整个单元格内容（作为最后的兜底）
                    # 注意：这里检查的是当前单元格是否已经被处理（通过检查 results 中是否包含 s 的子串）
                    already_extracted = any(s in r or r in s for r in results)
                    if not already_extracted and len(s) > 50 and re.search(r'(已采取的控制措施|人防措施)', s):
                        # 只返回包含关键词的部分
                        parts = s.split('\n')
                        relevant_parts = [p for p in parts if re.search(r'(已采取的控制措施|人防措施|专职|巡护)', p)]
                        if relevant_parts:
                            full_text = '\n'.join(relevant_parts).strip()
                            if len(full_text) > 20:
                                results.append(full_text)
                                if debug:
                                    debug_info.append(f"    ✓ 策略5提取成功: {len(full_text)}字符")

        # 如果所有策略都失败，尝试从整个表格中提取所有包含关键词的单元格
        if not results:
            if debug:
                debug_info.append("\n所有策略都未提取到结果，尝试提取所有相关单元格...")
            for tbl in doc.tables:
                tbl_text = "\n".join("  ".join(_cell_text(c) for c in row.cells) for row in tbl.rows)
                if re.search(r'(风险评价结果|风险评价结果表|风险评价)', tbl_text):
                    for row in tbl.rows:
                        for cell in row.cells:
                            s_raw = _cell_text(cell)
                            s = _norm(s_raw)
                            # 如果单元格较长且包含关键词，直接返回
                            if len(s) > 30 and re.search(r'(已采取的控制措施|人防措施|专职|巡护|控制措施)', s):
                                results.append(s)
                                if debug:
                                    debug_info.append(f"    兜底策略提取: {len(s)}字符的单元格内容")

        # 去重与清理
        uniq, seen = [], set()
        for t in results:
            normalized = re.sub(r'\s+', ' ', t).strip()
            if normalized and len(normalized) > 5 and normalized not in seen:
                seen.add(normalized)
                uniq.append(normalized)
        
        if debug:
            debug_info.append(f"\n最终提取结果: {len(uniq)}条")
            for idx, item in enumerate(uniq, 1):
                debug_info.append(f"  {idx}. {len(item)}字符: {item[:100]}...")
            print("\n".join(debug_info))
        
        return uniq
    except Exception as e:
        import traceback
        # 调试：打印错误信息
        error_msg = f"提取控制措施时出错: {e}\n{traceback.format_exc()}"
        if debug:
            print(error_msg)
        return []

def pick_district_leaders_from_measures(measures):
    """从控制措施文本中提取专职区（段）长姓名"""
    # 匹配模式：匹配标签后直接提取2-4个中文字符（最常见的中文姓名长度）
    # 这样可以避免包含后面的文字
    pat = re.compile(
        r'(?:专职\s*区[（(]段[）)]长(?:姓名)?|专职\s*区长|专职\s*段长|专职区段长)\s*[:：]\s*([一-龥·]{2,4})(?:\s|[\n\r]|[,，。;；:：]|[^\s\n\r,，。;；:：一-龥·]|$)',
        re.IGNORECASE
    )
    
    # 备用模式：如果精确匹配失败，使用更宽松的匹配
    pat_loose = re.compile(
        r'(?:专职\s*区[（(]段[）)]长(?:姓名)?|专职\s*区长|专职\s*段长|专职区段长)\s*[:：]\s*([一-龥·]{2,10})',
        re.IGNORECASE
    )
    
    names, seen = [], set()
    for s in measures:
        # 标准化括号
        s_norm = s.replace("(", "（").replace(")", "）")
        
        # 先尝试精确匹配（2-4个中文字符）
        matched = False
        for match in pat.finditer(s_norm):
            n = match.group(1).strip() if match.group(1) else ""
            if n and 2 <= len(n) <= 4:
                if n not in seen:
                    seen.add(n)
                    names.append(n)
                    matched = True
                    break
        
        # 如果精确匹配失败，使用宽松匹配并提取前2-4个字符
        if not matched:
            for match in pat_loose.finditer(s_norm):
                n = match.group(1).strip() if match.group(1) else ""
                if n:
                    # 只取中文字符
                    chinese_chars = re.findall(r'[一-龥·]', n)
                    if len(chinese_chars) >= 2:
                        # 常见的中文姓名长度为2-3个字符，优先匹配2-3个字符
                        # 常见动词/助词列表（如果第3个或第4个字符是这些，说明不是姓名的一部分）
                        common_verbs = ['该', '的', '是', '为', '有', '会', '要', '可', '能', '应', '需', '须', '每', '隔', '将', '在', '于', '对', '从', '向']
                        
                        # 策略1：优先尝试3个字符（三字名）
                        if len(chinese_chars) >= 3:
                            third_char = chinese_chars[2]
                            # 如果第3个字符不是动词，取前3个
                            if third_char not in common_verbs:
                                name_candidate = ''.join(chinese_chars[:3])
                                if name_candidate not in seen:
                                    seen.add(name_candidate)
                                    names.append(name_candidate)
                                    break
                        
                        # 策略2：如果3个字符不满足，或第3个字符是动词，尝试2个字符
                        if len(chinese_chars) >= 2:
                            second_char = chinese_chars[1]
                            # 如果第2个字符不是动词，取前2个
                            if second_char not in common_verbs:
                                name_candidate = ''.join(chinese_chars[:2])
                                if name_candidate not in seen:
                                    seen.add(name_candidate)
                                    names.append(name_candidate)
                                    break
                        
                        # 策略3：如果提取了4个或更多字符，检查第4个字符
                        if len(chinese_chars) >= 4:
                            fourth_char = chinese_chars[3]
                            if fourth_char in common_verbs:
                                # 第4个字符是动词，肯定不是姓名，取前3个或前2个
                                if len(chinese_chars) >= 3:
                                    name_candidate = ''.join(chinese_chars[:3])
                                    if name_candidate not in seen:
                                        seen.add(name_candidate)
                                        names.append(name_candidate)
                                        break
                                else:
                                    name_candidate = ''.join(chinese_chars[:2])
                                    if name_candidate not in seen:
                                        seen.add(name_candidate)
                                        names.append(name_candidate)
                                        break
                            else:
                                # 第4个字符不是动词，可能是复姓，取前4个
                                name_candidate = ''.join(chinese_chars[:4])
                                if name_candidate not in seen:
                                    seen.add(name_candidate)
                                    names.append(name_candidate)
                                    break
    
    return names

def _is_valid_docx_path(p: Path) -> bool:
    return p.suffix.lower() == ".docx" and not p.name.startswith("~$") and p.exists()


def _auto_discover_docx() -> Path | None:
    cwd = Path.cwd()
    candidates = []
    wm = cwd / "word-master"
    if wm.exists():
        for p in wm.rglob("*.docx"):
            if _is_valid_docx_path(p):
                candidates.append(p)
    for p in cwd.glob("*.docx"):
        if _is_valid_docx_path(p):
            candidates.append(p)
    if not candidates:
        for p in cwd.rglob("*.docx"):
            if _is_valid_docx_path(p):
                candidates.append(p)
                break
    return candidates[0] if candidates else None


# ========================== 4. 一致性检查（无模型） ==========================
def _normalize_text(s: str) -> str:
    return re.sub(r"\s+", "", (s or "")).strip()


def check_cross_references(sections, doc_path=None):
    leaders = extract_section_leaders(sections)
    locations = extract_locations(sections)
    # 如果提供了doc_path，使用优化后的方法提取编号
    if doc_path:
        hca_numbers_dict = extract_hca_numbers(doc_path)
        cover_number = hca_numbers_dict.get("封面编号", "未找到")
        table_number = hca_numbers_dict.get("高后果区基本信息表编号", "未找到")
    else:
        # 兼容旧方法：从sections中提取（降级方案）
        hca_numbers_list = extract_hca_numbers_legacy(sections)
        cover_number = hca_numbers_list[0] if hca_numbers_list else "未找到"
        table_number = hca_numbers_list[1] if len(hca_numbers_list) > 1 else "未找到"
        hca_numbers_dict = {
            "封面编号": cover_number,
            "高后果区基本信息表编号": table_number
        }

    # 区段长一致性
    pre_names = {t.split("：", 1)[-1] for t in leaders.get("from_prevention_measures") or []}
    base_names = {t.split("：", 1)[-1] for t in leaders.get("from_basic_info_table") or []}
    leader_conflict = False
    leader_detail = {
        "prevention": leaders.get("from_prevention_measures") or [],
        "basic_info": leaders.get("from_basic_info_table") or [],
        "risk_table": leaders.get("from_risk_table") or []
    }
    if pre_names and base_names and pre_names.isdisjoint(base_names):
        leader_conflict = True

    # 位置信息一致性（去除前缀，仅比对具体位置文本）
    def extract_loc_text(items):
        vals = []
        for x in items or []:
            parts = x.split("：", 1)
            vals.append(_normalize_text(parts[1] if len(parts) == 2 else x))
        return set(vals)

    loc_basic = extract_loc_text(locations.get("from_basic_info"))
    loc_risk = extract_loc_text(locations.get("from_risk_table"))
    location_conflict = False
    location_detail = {
        "basic_info": locations.get("from_basic_info") or [],
        "risk_table": locations.get("from_risk_table") or []
    }
    if loc_basic and loc_risk and loc_basic.isdisjoint(loc_risk):
        location_conflict = True

    # 高后果区编号一致性：判断封面编号和高后果区基本信息表编号是否一致
    # 排除都未找到的情况
    if cover_number == "未找到" and table_number == "未找到":
        hca_conflict = False  # 数据不足，不算不一致
        enough_hca = False
    else:
        enough_hca = True
        hca_conflict = (cover_number != table_number)

    return {
        "leaders": leaders,
        "locations": locations,
        "hca_numbers": hca_numbers_dict,
        "leader_inconsistent": leader_conflict,
        "location_inconsistent": location_conflict,
        "hca_number_inconsistent": hca_conflict,
        "details": {
            "leaders": leader_detail,
            "locations": location_detail,
        }
    }


def build_cross_reference_from_docx(doc_path: str, debug=False) -> str:
    """
    针对单个 .docx 返回上下文一致性检查的人类可读文本（不写文件，不调用 VLM）。
    """
    if not os.path.exists(doc_path):
        return f"❌ 找不到文档：{doc_path}"

    sections = parse_docx_sections(str(doc_path))
    rep = check_cross_references(sections, doc_path=str(doc_path))
    control_measures = extract_control_measures(str(doc_path), debug=debug)
    # 从控制措施中提取专职区（段）长姓名
    district_leaders_from_measures = pick_district_leaders_from_measures(control_measures)

    lines = []
    # 提取统计
    lines.append("上下文一致性检查：")
    # 先输出内容
    lines.append("内容：")
    lines.append("- 区段长信息：")
    lines.append("  人防措施：" + ("；".join(rep['leaders']['from_prevention_measures']) or "无"))
    lines.append("  基本信息表：" + ("；".join(rep['leaders']['from_basic_info_table']) or "无"))
    # lines.append("  风险评价表：" + ("；".join(rep['leaders'].get('from_risk_table', [])) or "无"))
    if district_leaders_from_measures:
        leaders_with_title = ["专职区（段）长：" + name for name in district_leaders_from_measures]
        lines.append("  风险评价表：" + "；".join(leaders_with_title))
    else:
        lines.append("  风险评价表：无")
    lines.append("- 位置信息：")
    lines.append("  基本信息表：" + ("；".join(rep['locations']['from_basic_info']) or "无"))
    lines.append("  风险评价表：" + ("；".join(rep['locations']['from_risk_table']) or "无"))
    lines.append("- 高后果区编号：")
    hca_dict = rep['hca_numbers']
    if isinstance(hca_dict, dict):
        lines.append("  封面编号：" + hca_dict.get("封面编号", "未找到"))
        lines.append("  高后果区基本信息表编号：" + hca_dict.get("高后果区基本信息表编号", "未找到"))
    else:
        # 兼容旧格式
        lines.append("  高后果区编号：" + ("，".join(rep['hca_numbers']) if rep['hca_numbers'] else "无"))
    lines.append("")

    # 一致性结论
    def verdict(flag: bool, enough: bool) -> str:
        if not enough:
            return "数据不足（无法对比）"
        return "不一致" if flag else "一致"

    enough_leader = bool(rep['leaders']['from_prevention_measures']) and bool(rep['leaders']['from_basic_info_table'])
    enough_location = bool(rep['locations']['from_basic_info']) and bool(rep['locations']['from_risk_table'])
    
    # 判断是否有足够的编号数据进行对比
    hca_dict = rep['hca_numbers']
    if isinstance(hca_dict, dict):
        cover_num = hca_dict.get("封面编号", "未找到")
        table_num = hca_dict.get("高后果区基本信息表编号", "未找到")
        enough_hca = (cover_num != "未找到" and table_num != "未找到")
    else:
        # 兼容旧格式
        enough_hca = len(rep['hca_numbers']) >= 1

    lines.append(f"结论：")
    lines.append(f"- 区段长信息：{verdict(rep['leader_inconsistent'], enough_leader)}")
    lines.append(f"- 位置信息：{verdict(rep['location_inconsistent'], enough_location)}")
    lines.append(f"- 高后果区编号：{verdict(rep['hca_number_inconsistent'], enough_hca)}")

    # 可选详情（只在不一致时给出两侧值，方便定位）
    if rep['leader_inconsistent'] and enough_leader:
        lines.append("\n[详情] 区段长信息：")
        lines.append("人防措施：" + "；".join(rep['details']['leaders']['prevention']))
        lines.append("基本信息表：" + "；".join(rep['details']['leaders']['basic_info']))
    if rep['location_inconsistent'] and enough_location:
        lines.append("\n[详情] 位置信息：")
        lines.append("基本信息表：" + "；".join(rep['details']['locations']['basic_info']))
        lines.append("风险评价表：" + "；".join(rep['details']['locations']['risk_table']))
    if rep['hca_number_inconsistent']:
        hca_dict = rep['hca_numbers']
        if isinstance(hca_dict, dict):
            lines.append("\n[详情] 高后果区编号：")
            lines.append("封面编号：" + hca_dict.get("封面编号", "未找到"))
            lines.append("高后果区基本信息表编号：" + hca_dict.get("高后果区基本信息表编号", "未找到"))
        else:
            # 兼容旧格式
            lines.append("\n[详情] 高后果区编号：" + "，".join(rep['hca_numbers']))

    return "\n".join(lines)


# ========================== 5. CLI ==========================
def parse_args():
    ap = argparse.ArgumentParser(description="上下文一致性检查（不调用大模型）")
    ap.add_argument("--input", required=False, default=None, help="输入 DOCX 文件路径；不提供则优先使用样例，否则自动查找")
    ap.add_argument("--debug", action="store_true", help="启用调试模式，显示详细的提取过程")
    return ap.parse_args()


def main():
    args = parse_args()
    chosen: Path | None
    if args.input:
        p = Path(args.input)
        chosen = p if _is_valid_docx_path(p) else None
    else:
        sample = Path(INPUT_DOCUMENT_PATH)
        if _is_valid_docx_path(sample):
            chosen = sample
        else:
            chosen = _auto_discover_docx()

    if not chosen:
        print("未找到可处理的 DOCX 文件，请使用 --input 指定文件路径。")
        return

    text = build_cross_reference_from_docx(str(chosen), debug=args.debug)
    print(text)


if __name__ == "__main__":
    main()