import os
import re
import json
import argparse
from pathlib import Path
from typing import List, Dict

SAMPLE_FILE_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/豫洛阳-兰郑长干线-CPY-0790-BFGDGS-ZZSYQFGS.docx"
SAMPLE_OUT_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/entry_route_result.txt"

ENTRY_ROUTE_ALIASES = [
    "入场线路图",
    "入场线路",
]
EVAC_ALIASES = [
    "逃生路线图",
    "逃生路线",
]
WINDOW = 2

try:
    from docx import Document
except ImportError:
    raise ImportError("请先 pip install python-docx")

def parse_args():
    ap = argparse.ArgumentParser(description="识别‘(1)入场线路图’到‘(2)逃生路线’之间的图片")
    ap.add_argument("--input", required=False, default=None, help="输入 DOCX 文件路径；不提供则优先使用样例，否则自动查找")
    ap.add_argument("--output_dir", default="入场线路图输出", help="图片输出目录")
    ap.add_argument("--out_path", default=None, help="结果写入该文本/JSON 文件（可选）")
    return ap.parse_args()

def _is_valid_docx_path(p: Path) -> bool:
    return p.suffix.lower() == ".docx" and not p.name.startswith("~$") and p.exists()

def _auto_discover_docx() -> Path|None:
    cwd = Path.cwd()
    candidates = []
    wm = cwd / "word-master"
    if wm.exists():
        for p in wm.rglob("*.docx"):
            if _is_valid_docx_path(p):
                candidates.append(p)
    for p in cwd.glob("*.docx"):
        if _is_valid_docx_path(p):
            candidates.append(p)
    if not candidates:
        for p in cwd.rglob("*.docx"):
            if _is_valid_docx_path(p):
                candidates.append(p)
                break
    return candidates[0] if candidates else None

def para_has_image(p):
    try:
        if p._element.xpath(
            './/*[local-name()="drawing" or local-name()="blip" or local-name()="imagedata" or local-name()="pic" or local-name()="pict" or local-name()="shape"]'
        ):
            return True
        for run in p.runs:
            el = run._element
            if el.xpath(
                './/*[local-name()="drawing" or local-name()="blip" or local-name()="imagedata" or local-name()="pic" or local-name()="pict" or local-name()="shape"]'
            ):
                return True
    except Exception:
        pass
    return False

def iter_block_items(doc):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    body = doc._element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield ("paragraph", Paragraph(child, doc), None)
        elif isinstance(child, CT_Tbl):
            tbl = Table(child, doc)
            for r_idx, row in enumerate(tbl.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p in cell.paragraphs:
                        yield ("table", p, (r_idx, c_idx))

def extract_images_between_numbered_captions(doc_path: str,
                                             start_aliases: list[str],
                                             end_aliases: list[str],
                                             output_dir: str) -> list:
    doc = Document(doc_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    doc_name = Path(doc_path).stem
    img_idx = 0
    results = []

    # 形如：(1) 或 1) 或 （1）
    def numbered_re(n: int, aliases: list[str]):
        alias_pat = "|".join(re.escape(a) for a in aliases)
        return re.compile(rf'^\s*[（(]?\s*{n}\s*[)）].*?(?:{alias_pat})')

    start_re = numbered_re(1, start_aliases)
    end_re   = numbered_re(2, end_aliases)

    in_range = False
    ended = False

    def save_from_runs_if_in_range(runs):
        nonlocal img_idx
        local = []
        for run in runs:
            drawing = run._element.xpath('.//a:blip')
            for blip in drawing:
                rEmbed = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rEmbed and rEmbed in doc.part.rels:
                    rel = doc.part.rels[rEmbed]
                    if "image" in rel.target_ref:
                        img_idx += 1
                        image_data = rel.target_part.blob
                        img_path = output_dir / f"{doc_name}_入场线路图_{img_idx}.png"
                        with open(img_path, 'wb') as f:
                            f.write(image_data)
                        local.append(str(img_path))
        return local

    # 扫描正文
    for para in doc.paragraphs:
        if ended:
            break
        text = para.text or ""
        if not in_range:
            if start_re.search(text):
                in_range = True
            continue
        if end_re.search(text):
            ended = True
            break
        results.extend(save_from_runs_if_in_range(para.runs))

    # 扫描表格
    if not ended:
        for table in doc.tables:
            if ended:
                break
            for row in table.rows:
                if ended:
                    break
                for cell in row.cells:
                    if ended:
                        break
                    cell_text = cell.text or ""
                    if not in_range:
                        if start_re.search(cell_text):
                            in_range = True
                        continue
                    if end_re.search(cell_text):
                        ended = True
                        break
                    for para in cell.paragraphs:
                        results.extend(save_from_runs_if_in_range(para.runs))

    return results


def build_vlm_messages_from_docx(docx_path: str,
                                 prompt_text: str = (
                                     "请基于入场线路图进行核查：\n"
                                     "1) 是否明确标注了从外部道路到施工现场的入场路线与关键路口；\n"
                                     "2) 是否包含必要的方位/地标/管线等辅助信息。\n"
                                     "请先给出‘是否符合要求：是/否’的结论；如不符合，请指出缺失或错误并给出改进建议。"
                                 ),
                                 output_dir: str = "入场线路图输出",
                                 first_only: bool = False) -> List[Dict]:
    """
    提取文档中“(1)入场线路图”到“(2)逃生路线”之间的图片，并构建与 VLM 兼容的消息结构：
    [
      {"role": "user", "content": [
          {"type": "image", "image": "..."},
          {"type": "text",  "text":  "...提示词..."}
      ]}
    ]
    """
    items = extract_images_between_numbered_captions(
        docx_path,
        ENTRY_ROUTE_ALIASES,
        EVAC_ALIASES,
        output_dir,
    ) or []

    if not items:
        return "无入场线路图，无需审核。"

    if first_only and items:
        items = [items[0]]

    content: List[Dict] = []
    for img_path in items:
        if img_path:
            content.append({"type": "image", "image": img_path})

    content.append({"type": "text", "text": prompt_text})

    return [
        {
            "role": "user",
            "content": content,
        }
    ]

def main():
    args = parse_args()
    chosen = None
    used_sample = False
    if args.input:
        chosen = Path(args.input)
    else:
        if SAMPLE_FILE_PATH and _is_valid_docx_path(Path(SAMPLE_FILE_PATH)):
            chosen = Path(SAMPLE_FILE_PATH)
            used_sample = True
        else:
            chosen = _auto_discover_docx()
    if not chosen or not Path(chosen).exists():
        print("未提供 --input，且未在当前目录或样例路径找到有效 DOCX（自动忽略临时文件 ~$.docx）。请用 --input 指定文件，例如：")
        print("python ruchangluxian.py --input C:/path/to/file.docx")
        return

    items = extract_images_between_numbered_captions(str(chosen), ENTRY_ROUTE_ALIASES, EVAC_ALIASES, args.output_dir)
    out_file = args.out_path if args.out_path else (SAMPLE_OUT_PATH if used_sample else None)

    if not items:
        # 与 content_consistency_check 风格一致：直接输出文本
        msg = "无入场线路图，无需审核。"
        print(msg)
        if out_file:
            try:
                Path(out_file).parent.mkdir(parents=True, exist_ok=True)
                with open(out_file, "w", encoding="utf-8") as f:
                    f.write(msg)
            except Exception as e:
                print(f"写入结果失败: {e}")
        return

    # 若有提取到图片，保留原有 JSON 输出
    result = {"image_path": (items[0] if items else None)}
    print(json.dumps(result, ensure_ascii=False, indent=2))
    if out_file:
        try:
            Path(out_file).parent.mkdir(parents=True, exist_ok=True)
            with open(out_file, "w", encoding="utf-8") as f:
                f.write(json.dumps(result, ensure_ascii=False, indent=2))
        except Exception as e:
            print(f"写入结果失败: {e}")

if __name__ == "__main__":
    main()
