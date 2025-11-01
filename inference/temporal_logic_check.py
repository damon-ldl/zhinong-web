import json
import re
import os
from pathlib import Path
import argparse
import datetime
from docx import Document

# ========================== 1. 配置区 ==========================
API_KEY = "BearerCubJNNWPCQjooxiEbLGB:IgSlCSEPDRuFXHJlXlMw"
URL = "https://spark-api-open.xf-yun.com/v2/chat/completions"
INPUT_DOCUMENT_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/豫洛阳-兰郑长干线-CPY-0790-BFGDGS-ZZSYQFGS.docx"
OUTPUT_JSON_PATH = r"C:/Users/cassi/Desktop/bisai/word-master/temporal_logic_check.json"

# 为与其它脚本对齐的样例路径常量
SAMPLE_FILE_PATH = INPUT_DOCUMENT_PATH
SAMPLE_OUT_PATH = OUTPUT_JSON_PATH


# ========================== 2. 功能函数区 ==========================


def extract_text_from_docx(docx_path):
    """提取Word文档中的纯文本内容"""
    try:
        doc = Document(docx_path)
        full_text = ""
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text += text + "\n"

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text += "  ".join(row_text) + "\n"

        # 🔧 自动修正跨行时间标识（关键修改点）
        # 例如 “识别时间\n2024年3月” → “识别时间：2024年3月”
        full_text = re.sub(r'(识别时间)\s*\n\s*(\d{4}年\d{1,2}月)', r'\1：\2', full_text)
        full_text = re.sub(r'(风险评价时间)\s*\n\s*(\d{4}年\d{1,2}月)', r'\1：\2', full_text)
        full_text = re.sub(r'(评价时间)\s*\n\s*(\d{4}年\d{1,2}月)', r'\1：\2', full_text)
        return full_text
    except Exception as e:
        print(f"Word文档文本提取失败：{str(e)}")
        return ""


def extract_date_from_text(date_text):
    """从文本中提取日期"""
    date_text = re.sub(r'[^\d年月日./\-\s]', '', date_text)
    date_patterns = [
        r'(\d{4})[年.-](\d{1,2})[月.-](\d{1,2})',
        r'(\d{4})[/.-](\d{1,2})[/.-](\d{1,2})',
        r'(\d{4})\s*年\s*(\d{1,2})\s*月',
        r'(\d{1,2})[月.-](\d{1,2})[日，]?\s*(\d{4})',
    ]
    for pattern in date_patterns:
        match = re.search(pattern, date_text)
        if match:
            try:
                groups = match.groups()
                if len(groups) == 3:
                    if len(groups[0]) == 4:
                        year, month, day = int(groups[0]), int(groups[1]), int(groups[2])
                    else:
                        month, day, year = int(groups[0]), int(groups[1]), int(groups[2])
                elif len(groups) == 2:
                    year, month, day = int(groups[0]), int(groups[1]), 1
                if 2000 <= year <= 2035 and 1 <= month <= 12 and 1 <= day <= 31:
                    return datetime.date(year, month, day)
            except:
                continue
    return None


def extract_time_information(text):
    """提取时间信息"""
    time_info = {
        'cover_dates': [],
        'identification_dates': [],
        'risk_assessment_dates': [],
        'other_dates': []
    }

    # 封面时间
    cover_patterns = [
        r'编制时间[：:\s]*([^\n]+)',
        r'制定时间[：:\s]*([^\n]+)',
        r'发布时间[：:\s]*([^\n]+)',
        r'版本时间[：:\s]*([^\n]+)'
    ]
    for pattern in cover_patterns:
        matches = re.finditer(pattern, text)
        for match in matches:
            date_text = match.group(1).strip()
            extracted_date = extract_date_from_text(date_text)
            if extracted_date:
                time_info['cover_dates'].append({
                    'raw_text': date_text,
                    'parsed_date': extracted_date,
                    'context': match.group(0),
                    'position': match.span()
                })

    # 识别时间
    identification_patterns = [
        r'识别时间[：:\s]*([^\n]*)',
        r'识别日期[：:\s]*([^\n]*)',
        r'完成识别.*?(\d{4}年\d{1,2}月)',
    ]
    for pattern in identification_patterns:
        matches = re.finditer(pattern, text)
        for match in matches:
            date_text = match.group(1).strip()
            extracted_date = extract_date_from_text(date_text)
            if extracted_date:
                time_info['identification_dates'].append({
                    'raw_text': date_text,
                    'parsed_date': extracted_date,
                    'context': match.group(0),
                    'position': match.span()
                })

    # 风险评价时间
    assessment_patterns = [
        r'风险评价时间[：:\s]*([^\n]*)',
        r'风险评价日期[：:\s]*([^\n]*)',
        r'评价时间[：:\s]*([^\n]*)',
        r'完成评价.*?(\d{4}年\d{1,2}月)',
    ]
    for pattern in assessment_patterns:
        matches = re.finditer(pattern, text)
        for match in matches:
            date_text = match.group(1).strip()
            extracted_date = extract_date_from_text(date_text)
            if extracted_date:
                time_info['risk_assessment_dates'].append({
                    'raw_text': date_text,
                    'parsed_date': extracted_date,
                    'context': match.group(0),
                    'position': match.span()
                })

    # ✅ 新增：只保留第一个风险评价时间，避免重复识别
    if len(time_info['risk_assessment_dates']) > 1:
        # print(f"⚠️ 检测到多个风险评价时间（共{len(time_info['risk_assessment_dates'])}个），仅保留第一个。")
        time_info['risk_assessment_dates'] = [time_info['risk_assessment_dates'][0]]

    return time_info


def build_time_snippets(text, time_info, context_window=120):
    """根据识别到的时间位置信息，从原文中切片生成上下文片段。"""
    snippets = []
    for key in ['cover_dates', 'identification_dates', 'risk_assessment_dates']:
        for item in time_info.get(key, []):
            start, end = item.get('position', (0, 0))
            s = max(0, start - context_window)
            e = min(len(text), end + context_window)
            snippet = text[s:e].strip()
            if snippet and snippet not in snippets:
                snippets.append(snippet)
    # 若没有定位到具体位置，则回退使用前 N 字符
    if not snippets and text:
        snippets.append(text[: max(300, min(800, len(text)))] )
    return snippets


def format_year_month(d):
    """将日期格式化为YYYY年M月。"""
    return f"{d.year}年{d.month}月"


def build_canonical_time_line(time_info):
    """构造规范单行文本：封面编制时间：xx，高后果区识别时间：xx，风险评价时间：xx"""
    cover = time_info.get('cover_dates', [])
    ident = time_info.get('identification_dates', [])
    assess = time_info.get('risk_assessment_dates', [])

    cover_str = format_year_month(cover[0]['parsed_date']) if cover else 'xx'
    ident_str = format_year_month(ident[0]['parsed_date']) if ident else 'xx'
    assess_str = format_year_month(assess[0]['parsed_date']) if assess else 'xx'

    return f"封面编制时间：{cover_str}，高后果区识别时间：{ident_str}，风险评价时间：{assess_str}"


def get_temporal_negative_prompt(canonical_line, context_text, variant="order_conflict"):
    """构造“时间逻辑对抗样本”提示词：包含规范单行与原始上下文片段两部分。"""
    variant_to_desc = {
        # 让“风险评价时间”未晚于“高后果区识别时间”，违背时间逻辑
        'order_conflict': (
            "任务：生成‘时间先后顺序不合理’的对抗样本。\n"
            "要求：\n"
            "1. 原始文本为规范单行：‘封面编制时间：YYYY年M月，高后果区识别时间：YYYY年M月，风险评价时间：YYYY年M月’。保持结构与字段名不变。\n"
            "2. 尽量不改年份（YYYY），优先仅改月份/日，使‘风险评价时间’未晚于‘高后果区识别时间’（可同月或更早）。\n"
            "3. 不修改非时间文本与实体名称。\n"
            "原始文本（规范单行）：\n---\n{canonical}\n---\n\n原始上下文片段：\n---\n{context}\n---\n"
        ),
        # 年份不一致（仅修改年份），指定不一致的两类时间
        'year_inconsistency_cover_identification': (
            "任务：生成‘年份不一致（封面编制时间 vs 高后果区识别时间）’的对抗样本。\n"
            "要求：\n"
            "1. 原始文本为规范单行，保持结构与字段名不变。\n"
            "2. 仅修改年份（YYYY），使‘封面编制时间’与‘高后果区识别时间’年份不同（如封面2023，识别2024）。\n"
            "3. 月份尽量保持不变；不修改非时间文本。\n"
            "4. 输出仅为修改后的文本。\n"
            "原始文本（规范单行）：\n---\n{canonical}\n---\n\n原始上下文片段：\n---\n{context}\n---\n"
        ),
        'year_inconsistency_cover_assessment': (
            "任务：生成‘年份不一致（封面编制时间 vs 风险评价时间）’的对抗样本。\n"
            "要求：\n"
            "1. 原始文本为规范单行，保持结构与字段名不变。\n"
            "2. 仅修改年份（YYYY），使‘封面编制时间’与‘风险评价时间’年份不同。\n"
            "3. 月份尽量保持不变；不修改非时间文本。\n"
            "4. 输出仅为修改后的文本。\n"
            "原始文本（规范单行）：\n---\n{canonical}\n---\n\n原始上下文片段：\n---\n{context}\n---\n"
        ),
        'year_inconsistency_identification_assessment': (
            "任务：生成‘年份不一致（高后果区识别时间 vs 风险评价时间）’的对抗样本。\n"
            "要求：\n"
            "1. 原始文本为规范单行，保持结构与字段名不变。\n"
            "2. 仅修改年份（YYYY），使‘高后果区识别时间’与‘风险评价时间’年份不同。\n"
            "3. 月份尽量保持不变；不修改非时间文本。\n"
            "4. 输出仅为修改后的文本。\n"
            "原始文本（规范单行）：\n---\n{canonical}\n---\n\n原始上下文片段：\n---\n{context}\n---\n"
        ),
    }
    head = variant_to_desc.get(variant, variant_to_desc['order_conflict'])
    # 用具体原始文本填充模板
    prompt_filled = head.format(canonical=canonical_line or "", context=context_text or "")
    return f"{prompt_filled}\n输出：仅返回修改后的文本，无需额外解释。"


def check_temporal_logic(time_info):
    """检查时间逻辑一致性"""
    logic_results = {
        'year_consistency': True,
        'temporal_order_correct': True,
        'is_correct': True,
        'issues': [],
        'time_analysis': {}
    }

    all_dates = {
        'cover': [d['parsed_date'] for d in time_info['cover_dates']],
        'identification': [d['parsed_date'] for d in time_info['identification_dates']],
        'risk_assessment': [d['parsed_date'] for d in time_info['risk_assessment_dates']]
    }

    logic_results['time_analysis'] = {
        'cover_dates_count': len(all_dates['cover']),
        'identification_dates_count': len(all_dates['identification']),
        'risk_assessment_dates_count': len(all_dates['risk_assessment'])
    }

    # 年份一致性（细分三类两两不一致）
    cover_years = set(d.year for d in all_dates['cover'])
    identification_years = set(d.year for d in all_dates['identification'])
    risk_years = set(d.year for d in all_dates['risk_assessment'])

    # 任意两类均存在时进行对比
    if cover_years and identification_years and cover_years != identification_years:
        logic_results['year_consistency'] = False
        logic_results['is_correct'] = False
        logic_results['issues'].append(
            f"年份不一致（封面编制时间 vs 高后果区识别时间）：封面{sorted(cover_years)}，识别{sorted(identification_years)}"
        )
    if cover_years and risk_years and cover_years != risk_years:
        logic_results['year_consistency'] = False
        logic_results['is_correct'] = False
        logic_results['issues'].append(
            f"年份不一致（封面编制时间 vs 风险评价时间）：封面{sorted(cover_years)}，评价{sorted(risk_years)}"
        )
    if identification_years and risk_years and identification_years != risk_years:
        logic_results['year_consistency'] = False
        logic_results['is_correct'] = False
        logic_results['issues'].append(
            f"年份不一致（高后果区识别时间 vs 风险评价时间）：识别{sorted(identification_years)}，评价{sorted(risk_years)}"
        )

    # 时间先后逻辑
    if all_dates['identification'] and all_dates['risk_assessment']:
        if min(all_dates['risk_assessment']) <= max(all_dates['identification']):
            logic_results['temporal_order_correct'] = False
            logic_results['is_correct'] = False
            logic_results['issues'].append("风险评价时间未晚于高后果区识别时间")

    # 时间合理性
    current_date = datetime.date.today()
    for date_type, dates in all_dates.items():
        for d in dates:
            if d > current_date:
                logic_results['issues'].append(f"{date_type}时间({d})为未来时间")
                logic_results['is_correct'] = False

    # 时间缺失
    if not all_dates['cover']:
        logic_results['issues'].append("缺少封面编制时间")
        logic_results['is_correct'] = False
    if not all_dates['identification']:
        logic_results['issues'].append("缺少高后果区识别时间")
        logic_results['is_correct'] = False
    if not all_dates['risk_assessment']:
        logic_results['issues'].append("缺少风险评价时间")
        logic_results['is_correct'] = False

    return logic_results


# ========================== 报告生成 ==========================
def build_temporal_report(time_info, logic) -> str:
    lines = []
    # 概要
    lines.append("时间逻辑检查：")
    lines.append("")

    # 规范单行
    def format_year_month(d):
        return f"{d.year}年{d.month}月"
    cover = time_info.get('cover_dates', [])
    ident = time_info.get('identification_dates', [])
    assess = time_info.get('risk_assessment_dates', [])
    cover_str = format_year_month(cover[0]['parsed_date']) if cover else 'xx'
    ident_str = format_year_month(ident[0]['parsed_date']) if ident else 'xx'
    assess_str = format_year_month(assess[0]['parsed_date']) if assess else 'xx'
    lines.append(f"规范汇总：封面编制时间：{cover_str}，高后果区识别时间：{ident_str}，风险评价时间：{assess_str}")
    lines.append("")

    # 明细
    lines.append("时间信息明细：")
    for key, items in time_info.items():
        if key not in ['cover_dates', 'identification_dates', 'risk_assessment_dates']:
            continue
        cn = {
            'cover_dates': '封面编制时间',
            'identification_dates': '高后果区识别时间',
            'risk_assessment_dates': '风险评价时间',
        }.get(key, key)
        lines.append(f"- {cn}（{len(items)} 个）：")
        for i, item in enumerate(items, 1):
            lines.append(f"  {i}. {item['context']} -> {item['parsed_date']}")
    lines.append("")

    # 结果
    if logic.get('issues'):
        lines.append("结论：存在问题")
        for idx, iss in enumerate(logic['issues'], 1):
            lines.append(f"- 问题{idx}：{iss}")
    else:
        lines.append("结论：无问题")

    return "\n".join(lines)


# ========================== 3. CLI 对齐（与其它脚本一致） ==========================
def _is_valid_docx_path(p: Path) -> bool:
    return p.suffix.lower() == ".docx" and not p.name.startswith("~$") and p.exists()


def _auto_discover_docx() -> Path | None:
    cwd = Path.cwd()
    candidates = []
    wm = cwd / "word-master"
    if wm.exists():
        for p in wm.rglob("*.docx"):
            if _is_valid_docx_path(p):
                candidates.append(p)
    for p in cwd.glob("*.docx"):
        if _is_valid_docx_path(p):
            candidates.append(p)
    if not candidates:
        for p in cwd.rglob("*.docx"):
            if _is_valid_docx_path(p):
                candidates.append(p)
                break
    return candidates[0] if candidates else None


def parse_args():
    ap = argparse.ArgumentParser(description="时间逻辑检查与对抗样本生成")
    ap.add_argument("--input", required=False, default=None, help="输入 DOCX 文件路径；不提供则优先使用样例，否则自动查找")
    ap.add_argument("--out_path", required=False, default=None, help="输出 JSON 文件路径（可选）")
    return ap.parse_args()


# ========================== 主程序 ==========================
if __name__ == '__main__':
    args = parse_args()
    chosen = None
    used_sample = False
    if args.input:
        p = Path(args.input)
        chosen = p if _is_valid_docx_path(p) else None
    else:
        sample = Path(SAMPLE_FILE_PATH)
        if _is_valid_docx_path(sample):
            chosen = sample
            used_sample = True
        else:
            chosen = _auto_discover_docx()

    if not chosen:
        print("未找到可处理的 DOCX 文件，请使用 --input 指定文件路径。")
        raise SystemExit(0)

    text = extract_text_from_docx(str(chosen))

    time_info = extract_time_information(text)

    logic = check_temporal_logic(time_info)

    # =============== 生成文本报告（不进行正负样本生成） ===============
    report = build_temporal_report(time_info, logic)

    # 持久化输出：优先使用 --out_path，否则样例路径；写入为纯文本
    out_file = args.out_path if getattr(args, 'out_path', None) else (SAMPLE_OUT_PATH if used_sample else None)
    if out_file:
        try:
            out_dir = os.path.dirname(out_file)
            if out_dir and not os.path.exists(out_dir):
                os.makedirs(out_dir)
            with open(out_file, "w", encoding="utf-8") as f:
                f.write(report)
        except Exception as e:
            print(f"保存报告失败：{e}")
    print(report)
